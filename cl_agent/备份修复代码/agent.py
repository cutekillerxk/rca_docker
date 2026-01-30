#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HDFS集群监控Agent（LangChain 1.0.7 + vLLM）
整合所有功能：日志读取、监控采集、Agent创建、文档导出
"""

import os
import re
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

# LangChain 导入
from langchain.agents import create_agent
from langchain_openai import ChatOpenAI

from langchain_tavily import TavilySearch
website_search = TavilySearch(max_results=2, tavily_api_key=os.getenv("TAVILY_API_KEY"))

# 导入配置
from .config import (
    VLLM_BASE_URL,
    VLLM_MODEL_PATH,
    THIRD_PARTY_API_BASE_URL,
    THIRD_PARTY_API_KEY
)

# 导入集群上下文配置（用于生成System Prompt）
from .cluster_context import generate_system_prompt

# 导入工具函数
from .tools import (
    get_cluster_logs,
    get_node_log,
    get_monitoring_metrics,
    search_logs_by_keyword,
    get_error_logs_summary,
    hadoop_auto_operation,
    execute_hadoop_command,
    generate_repair_plan
)

# ==================== LLM 配置 ====================

def create_llm(model_name: str = "qwen-8b") -> ChatOpenAI:
    """
    根据模型名称创建LLM实例
    
    Args:
        model_name: 模型名称，可选值：
            - "qwen-8b": Qwen-8B (vLLM本地部署)
            - "gpt-4o": GPT-4o (第三方API)
            - "deepseek-r1": DeepSeek-R1 (第三方API)
    
    Returns:
        ChatOpenAI 实例
    """
    # 模型配置字典
    model_configs = {
        "qwen-8b": {
            "base_url": VLLM_BASE_URL,
            "api_key": "not-needed",
            "model": VLLM_MODEL_PATH,
            "timeout": 120,
            "max_tokens": 4096,
        },
        "gpt-4o": {
            "base_url": THIRD_PARTY_API_BASE_URL,
            "api_key": THIRD_PARTY_API_KEY,
            "model": "gpt-4o",  # 实际模型名根据第三方API调整
            "timeout": 60,
            "max_tokens": 4096,
        },
        "deepseek-r1": {
            "base_url": THIRD_PARTY_API_BASE_URL,
            "api_key": THIRD_PARTY_API_KEY,
            "model": "DeepSeek-V3.2",  # 实际模型名根据第三方API调整
            "timeout": 120,
            "max_tokens": 4096,
        }
    }
    
    # 获取模型配置，如果不存在则使用默认配置（qwen-8b）
    config = model_configs.get(model_name, model_configs["qwen-8b"])
    
    # 检查第三方API配置（gpt-4o 和 deepseek-r1 需要）
    if model_name in ["gpt-4o", "deepseek-r1"]:
        if not config["base_url"] or not config["api_key"]:
            raise ValueError(
                f"模型 {model_name} 需要配置 API_BASE_URL 和 API_KEY 环境变量。\n"
                f"请在 .env 文件中设置：\n"
                f"  API_BASE_URL=你的第三方API地址\n"
                f"  API_KEY=你的API密钥"
            )
    
    # 调试信息：打印模型配置
    print(f"[DEBUG] 创建LLM实例 - 模型: {model_name}")
    print(f"[DEBUG]   - base_url: {config['base_url']}")
    print(f"[DEBUG]   - model: {config['model']}")
    # 安全显示 API key（只显示后4位）
    if config['api_key'] == "not-needed":
        api_key_display = "not-needed (vLLM本地部署)"
    elif config['api_key']:
        api_key_display = f"已设置 (***{config['api_key'][-4:]})" if len(config['api_key']) > 4 else "已设置 (***)"
    else:
        api_key_display = "未设置"
    print(f"[DEBUG]   - api_key: {api_key_display}")
    print(f"[DEBUG]   - timeout: {config['timeout']}s")
    print(f"[DEBUG]   - max_tokens: {config['max_tokens']}")
    
    llm = ChatOpenAI(
        base_url=config["base_url"],
        api_key=config["api_key"],
        model=config["model"],
        temperature=0,
        max_tokens=config["max_tokens"],
        timeout=config["timeout"],
        max_retries=2,
    )
    
    # 调试信息：确认LLM实例创建成功
    print(f"[DEBUG] ✅ LLM实例创建成功 - 模型: {model_name} (实际模型名: {config['model']})")
    
    return llm

# ==================== Agent创建 ====================

def create_agent_instance(model_name: str = "qwen-8b"):
    """
    根据模型名称创建Agent实例
    
    Args:
        model_name: 模型名称，可选值：
            - "qwen-8b": Qwen-8B (vLLM本地部署)
            - "gpt-4o": GPT-4o (第三方API)
            - "deepseek-r1": DeepSeek-R1 (第三方API)
    
    Returns:
        Agent 实例
    """
    print(f"[DEBUG] 开始创建Agent实例 - 模型: {model_name}")
    
    # 导入知识库模块（暂时注释，不使用知识库）
    # try:
    #     from .knowledge_base import search_operation_knowledge
    # except ImportError:
    #     from lc_agent.knowledge_base import search_operation_knowledge
    
    # 定义操作知识库检索工具（暂时注释，不使用知识库）
    # @tool("search_operation_knowledge", description="检索Hadoop集群操作知识库，获取操作命令格式、示例和说明。如果不确定如何执行某个操作，请先使用此工具检索。")
    # def search_operation_knowledge_tool(query: str) -> str:
    #     """检索操作知识库"""
    #     return search_operation_knowledge(query)
    
    llm = create_llm(model_name)
    tools = [get_cluster_logs, get_node_log, get_monitoring_metrics, website_search, hadoop_auto_operation, execute_hadoop_command, generate_repair_plan]

    # 使用cluster_context模块生成System Prompt
    # 包含完整的集群配置信息、命令格式模板、工作流程等
    system_prompt = generate_system_prompt()
    
    # 追加工具使用说明（补充cluster_context中未包含的工具相关信息）
    system_prompt += """

## 可用工具及使用场景

### 信息收集类
1. **get_cluster_logs** - 获取集群所有节点的最新日志
   - 场景：开始诊断时，了解各节点状态
   - 注意：按思考检查点逐个分析节点日志，最后汇总

2. **get_node_log(node_name)** - 获取指定节点的日志
   - 场景：深入分析某个特定节点

3. **get_monitoring_metrics** - 获取JMX监控指标
   - 场景：了解集群的量化状态（存活节点数、存储使用等）

### 命令执行类
4. **execute_hadoop_command(command_args)** - 执行Hadoop查询命令
   - 参数格式：["hdfs", "dfsadmin", "-report"]
   - 场景：执行hdfs、hadoop等命令获取信息

5. **hadoop_auto_operation(operation, container)** - 启动/停止/重启Hadoop服务
   - operation: "start" | "stop" | "restart"
   - container: "namenode" | "datanode1" | "datanode2" | None(整个集群)
   - 场景：管理Hadoop服务生命周期

### 计划生成类
6. **generate_repair_plan(fault_type, diagnosis_info, affected_nodes)** - 生成修复计划
   - 支持的故障类型：datanode_down, cluster_id_mismatch, namenode_safemode, datanode_disk_full, namenode_down, multiple_datanodes_down
   - 场景：诊断完成后，制定修复方案

### 网络搜索
7. **website_search** - 搜索Hadoop相关技术文档
   - 场景：遇到不常见问题时，搜索解决方案

## 修复操作流程（重要）

1. **诊断问题** - 使用 get_cluster_logs、get_monitoring_metrics、execute_hadoop_command
2. **生成修复计划** - 使用 generate_repair_plan 工具，输出JSON格式计划
3. **展示计划** - 将计划展示给用户，等待确认
4. **执行修复** - 用户确认后，使用 hadoop_auto_operation 或 execute_hadoop_command
5. **验证结果** - 检查集群状态，确认修复成功

请用专业、清晰的语言回答。"""
    
    print(f"[DEBUG] 正在创建Agent（使用 {model_name} 模型）...")
    agent = create_agent(
        model=llm,
        tools=tools,
        system_prompt=system_prompt
    )
    
    print(f"[DEBUG] ✅ Agent实例创建成功 - 模型: {model_name}")
    
    return agent

# ==================== 文档导出功能 ====================

def parse_markdown_to_word_paragraphs(doc, text: str):
    """
    将Markdown格式的文本解析并添加到Word文档中
    
    Args:
        doc: python-docx Document对象
        text: Markdown格式的文本
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    lines = text.split('\n')
    i = 0
    in_list = False
    list_level = 0
    current_para_lines = []  # 用于合并多行段落
    
    while i < len(lines):
        line = lines[i].rstrip()  # 只移除右侧空格，保留左侧缩进信息
        
        # 处理分隔线（---或===）
        if line.strip().startswith('---') or (line.strip().startswith('===') and len(line.strip()) >= 3):
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if in_list:
                in_list = False
                list_level = 0
            doc.add_paragraph()  # 空行作为分隔
            i += 1
            continue
        
        # 处理标题（###、####等）
        if line.strip().startswith('#'):
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if in_list:
                in_list = False
                list_level = 0
            
            # 计算标题级别
            stripped = line.strip()
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            
            # 提取标题文本（移除#号和空格）
            title_text = stripped[level:].strip()
            
            # 移除Markdown格式标记（**、`等），但保留文本
            title_text = re.sub(r'\*\*(.*?)\*\*', r'\1', title_text)  # 移除加粗，保留文本
            title_text = re.sub(r'`(.*?)`', r'\1', title_text)  # 移除代码标记，保留文本
            
            # 添加标题（Word标题级别：1-9，我们限制在1-3）
            heading_level = min(level, 3)
            doc.add_heading(title_text, heading_level)
            i += 1
            continue
        
        # 处理有序列表（1.、2.等）
        ordered_match = re.match(r'^(\d+)[\.、]\s+(.+)$', line.strip())
        if ordered_match:
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if not in_list:
                in_list = True
                list_level = 0
            
            item_text = ordered_match.group(2)
            # 处理列表项中的格式（保留格式，不转换为纯文本）
            para = doc.add_paragraph(style='List Number')
            _add_formatted_text(para, item_text)
            i += 1
            continue
        
        # 处理无序列表（-、*等，但排除分隔线）
        if line.strip() and (line.strip().startswith('- ') or line.strip().startswith('* ')):
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if not in_list:
                in_list = True
                list_level = 0
            
            # 提取列表项文本（移除 - 或 * 和空格）
            item_text = line.strip()[2:].strip()
            # 处理列表项中的格式（保留格式）
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, item_text)
            i += 1
            continue
        
        # 处理空行
        if not line.strip():
            if current_para_lines:
                # 处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if in_list:
                in_list = False
                list_level = 0
            # 添加空行
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理普通文本行
        if in_list:
            in_list = False
            list_level = 0
        
        # 累积到当前段落（处理多行段落）
        current_para_lines.append(line.strip())
        i += 1
    
    # 处理最后剩余的段落
    if current_para_lines:
        para = doc.add_paragraph()
        _add_formatted_text(para, ' '.join(current_para_lines))


def _process_inline_formatting(text: str) -> str:
    """处理行内格式标记，返回纯文本"""
    # 移除加粗标记
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # 移除代码标记
    text = re.sub(r'`(.*?)`', r'\1', text)
    # 移除斜体标记
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text


def _add_formatted_text(para, text: str):
    """
    将包含Markdown格式的文本添加到段落中，保留格式
    
    Args:
        para: python-docx Paragraph对象
        text: 包含Markdown格式的文本
    """
    from docx.shared import Pt, RGBColor
    
    if not text:
        return
    
    # 使用正则表达式匹配所有格式标记（加粗、代码、斜体）
    # 按顺序处理，避免嵌套问题
    pattern = r'(\*\*.*?\*\*|`.*?`|\*.*?\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # 处理加粗文本 **text**
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            bold_text = part[2:-2]  # 移除 **
            # 直接输出加粗文本，不再递归处理内部格式
            # 避免无限递归问题
            run = para.add_run(bold_text)
            run.bold = True
        
        # 处理代码文本 `text`
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            code_text = part[1:-1]  # 移除 `
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        # 处理斜体文本 *text*（但不是加粗）
        elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
            italic_text = part[1:-1]  # 移除 *
            run = para.add_run(italic_text)
            run.italic = True
        
        # 普通文本 - 直接输出，不再递归
        # 注意：如果文本中有不成对的格式标记（如单个 * 或 **），直接作为普通文本处理
        # 避免无限递归
        else:
            para.add_run(part)


def export_to_word(analysis_result: str, output_path: str) -> str:
    """
    将Agent的分析结果导出为Word文档
    
    Args:
        analysis_result: Agent返回的分析结果文本
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 标题
        title = doc.add_heading('HDFS集群问题诊断报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 时间戳
        timestamp = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
        time_para = doc.add_paragraph(f'生成时间: {timestamp}')
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # 空行
        
        # 执行摘要
        doc.add_heading('执行摘要', 1)
        summary_para = doc.add_paragraph('本报告基于HDFS集群监控数据和日志分析生成，包含监控指标、问题诊断、原因分析和解决方案。')
        
        doc.add_paragraph()  # 空行
        
        # 添加集群监控指标表格
        doc.add_heading('集群监控指标', 1)
        
        try:
            from .monitor_collector import collect_all_metrics
            metrics_data = collect_all_metrics()
            
            # NameNode 指标表格
            doc.add_heading('NameNode 监控指标', 2)
            nn_table = doc.add_table(rows=1, cols=3)
            nn_table.style = 'Light Grid Accent 1'
            
            # 表头
            header_cells = nn_table.rows[0].cells
            header_cells[0].text = '指标名称'
            header_cells[1].text = '指标值'
            header_cells[2].text = '状态'
            
            # 设置表头样式
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # 添加NameNode指标数据
            if metrics_data["namenode"].get("status") != "error":
                for key, metric in metrics_data["namenode"].get("metrics", {}).items():
                    row_cells = nn_table.add_row().cells
                    row_cells[0].text = metric['name']
                    row_cells[1].text = str(metric['value'])
                    status_text = "正常" if metric["status"] == "normal" else "异常"
                    row_cells[2].text = status_text
                    
                    # 设置状态列颜色
                    if metric["status"] == "normal":
                        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                    else:
                        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
            else:
                row_cells = nn_table.add_row().cells
                row_cells[0].text = "错误"
                row_cells[1].text = metrics_data['namenode'].get('error', '未知错误')
                row_cells[2].text = "异常"
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph()  # 空行
            
            # DataNode 指标表格
            doc.add_heading('DataNode 监控指标', 2)
            for node_name, node_data in metrics_data["datanodes"].items():
                doc.add_heading(f'DataNode {node_name.upper()}', 3)
                dn_table = doc.add_table(rows=1, cols=3)
                dn_table.style = 'Light Grid Accent 1'
                
                # 表头
                header_cells = dn_table.rows[0].cells
                header_cells[0].text = '指标名称'
                header_cells[1].text = '指标值'
                header_cells[2].text = '状态'
                
                # 设置表头样式
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                
                # 添加DataNode指标数据
                if node_data.get("status") != "error":
                    for key, metric in node_data.get("metrics", {}).items():
                        row_cells = dn_table.add_row().cells
                        # 移除指标名称中的节点前缀，使表格更简洁
                        metric_name = metric['name'].replace(f'DataNode {node_name} ', '')
                        row_cells[0].text = metric_name
                        row_cells[1].text = str(metric['value'])
                        status_text = "正常" if metric["status"] == "normal" else "异常"
                        row_cells[2].text = status_text
                        
                        # 设置状态列颜色
                        if metric["status"] == "normal":
                            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                else:
                    row_cells = dn_table.add_row().cells
                    row_cells[0].text = "错误"
                    row_cells[1].text = node_data.get('error', '未知错误')
                    row_cells[2].text = "异常"
                    row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                
                doc.add_paragraph()  # 空行
            
            # 添加监控数据时间戳
            doc.add_paragraph(f'监控数据采集时间: {metrics_data["timestamp"]}')
            doc.add_paragraph()  # 空行
            
        except Exception as e:
            doc.add_paragraph(f'⚠️ 获取监控指标失败: {str(e)}')
            doc.add_paragraph()  # 空行
        
        doc.add_paragraph()  # 空行
        
        # 问题分析
        doc.add_heading('问题分析', 1)
        
        # 使用Markdown解析函数处理分析结果
        parse_markdown_to_word_paragraphs(doc, analysis_result)
        
        doc.add_paragraph()  # 空行
        
        # 解决方案
        doc.add_heading('解决方案', 1)
        solution_para = doc.add_paragraph('根据上述分析，建议采取以下措施：')
        solution_para.add_run('\n\n请参考问题分析部分的详细说明。')
        
        doc.add_paragraph()  # 空行
        
        # 预防措施
        doc.add_heading('预防措施', 1)
        prevention_para = doc.add_paragraph('为避免类似问题再次发生，建议：')
        prevention_para.add_run('\n1. 定期检查集群监控指标\n2. 及时处理异常日志\n3. 建立完善的监控告警机制')
        
        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"导出Word文档失败: {str(e)}")


def export_to_pdf(analysis_result: str, output_path: str) -> str:
    """
    将Agent的分析结果导出为PDF文档
    
    Args:
        analysis_result: Agent返回的分析结果文本
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import black
        
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=black,
            spaceAfter=30,
            alignment=1  # 居中
        )
        
        # 小标题样式
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=black,
            spaceAfter=12,
            spaceBefore=12
        )
        
        # 正文样式
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=6
        )
        
        # 标题
        story.append(Paragraph('HDFS集群问题诊断报告', title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 时间戳
        timestamp = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
        story.append(Paragraph(f'生成时间: {timestamp}', styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # 执行摘要
        story.append(Paragraph('执行摘要', heading_style))
        story.append(Paragraph('本报告基于HDFS集群监控数据和日志分析生成，包含监控指标、问题诊断、原因分析和解决方案。', normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 添加集群监控指标表格
        story.append(Paragraph('集群监控指标', heading_style))
        
        try:
            from .monitor_collector import collect_all_metrics
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            metrics_data = collect_all_metrics()
            
            # NameNode 指标表格
            story.append(Paragraph('NameNode 监控指标', heading_style))
            
            if metrics_data["namenode"].get("status") != "error":
                # 准备表格数据和样式
                nn_data = [['指标名称', '指标值', '状态']]
                table_styles = [
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ]
                
                row_idx = 1
                for key, metric in metrics_data["namenode"].get("metrics", {}).items():
                    status_text = "正常" if metric["status"] == "normal" else "异常"
                    nn_data.append([metric['name'], str(metric['value']), status_text])
                    
                    # 设置状态列颜色
                    status_color = colors.green if metric["status"] == "normal" else colors.red
                    table_styles.append(('TEXTCOLOR', (2, row_idx), (2, row_idx), status_color))
                    row_idx += 1
                
                # 创建表格
                nn_table = Table(nn_data, colWidths=[2.5*inch, 2*inch, 1*inch])
                nn_table.setStyle(TableStyle(table_styles))
                
                story.append(nn_table)
                story.append(Spacer(1, 0.2*inch))
            else:
                story.append(Paragraph(f'错误: {metrics_data["namenode"].get("error", "未知错误")}', normal_style))
                story.append(Spacer(1, 0.2*inch))
            
            # DataNode 指标表格
            story.append(Paragraph('DataNode 监控指标', heading_style))
            
            for node_name, node_data in metrics_data["datanodes"].items():
                story.append(Paragraph(f'DataNode {node_name.upper()}', heading_style))
                
                if node_data.get("status") != "error":
                    # 准备表格数据和样式
                    dn_data = [['指标名称', '指标值', '状态']]
                    table_styles = [
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    ]
                    
                    row_idx = 1
                    for key, metric in node_data.get("metrics", {}).items():
                        metric_name = metric['name'].replace(f'DataNode {node_name} ', '')
                        status_text = "正常" if metric["status"] == "normal" else "异常"
                        dn_data.append([metric_name, str(metric['value']), status_text])
                        
                        # 设置状态列颜色
                        status_color = colors.green if metric["status"] == "normal" else colors.red
                        table_styles.append(('TEXTCOLOR', (2, row_idx), (2, row_idx), status_color))
                        row_idx += 1
                    
                    # 创建表格
                    dn_table = Table(dn_data, colWidths=[2.5*inch, 2*inch, 1*inch])
                    dn_table.setStyle(TableStyle(table_styles))
                    
                    story.append(dn_table)
                else:
                    story.append(Paragraph(f'错误: {node_data.get("error", "未知错误")}', normal_style))
                
                story.append(Spacer(1, 0.2*inch))
            
            # 添加监控数据时间戳
            story.append(Paragraph(f'监控数据采集时间: {metrics_data["timestamp"]}', normal_style))
            story.append(Spacer(1, 0.2*inch))
            
        except Exception as e:
            story.append(Paragraph(f'⚠️ 获取监控指标失败: {str(e)}', normal_style))
            story.append(Spacer(1, 0.2*inch))
        
        # 问题分析
        story.append(Paragraph('问题分析', heading_style))
        
        # 将分析结果分段处理
        paragraphs = analysis_result.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 转义HTML特殊字符
                para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # 检查是否是标题
                if re.match(r'^[0-9]+[\.、]', para_text) or any(keyword in para_text for keyword in ['问题', '原因', '解决', '建议']):
                    if len(para_text) < 50:
                        story.append(Paragraph(para_text, heading_style))
                    else:
                        story.append(Paragraph(para_text, normal_style))
                else:
                    story.append(Paragraph(para_text, normal_style))
        
        story.append(Spacer(1, 0.2*inch))
         
        # 解决方案
        story.append(Paragraph('解决方案', heading_style))
        story.append(Paragraph('根据上述分析，建议采取以下措施：', normal_style))
        story.append(Paragraph('请参考问题分析部分的详细说明。', normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 预防措施
        story.append(Paragraph('预防措施', heading_style))
        story.append(Paragraph('为避免类似问题再次发生，建议：', normal_style))
        story.append(Paragraph('1. 定期检查集群监控指标', normal_style))
        story.append(Paragraph('2. 及时处理异常日志', normal_style))
        story.append(Paragraph('3. 建立完善的监控告警机制', normal_style))
        
        # 生成PDF
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.build(story)
        
        return output_path
    except ImportError:
        raise ImportError("需要安装 reportlab: pip install reportlab")
    except Exception as e:
        raise Exception(f"导出PDF文档失败: {str(e)}")

# ==================== 主函数（用于测试） ====================

if __name__ == "__main__":
    # 测试Agent创建
    print("[INFO] 正在创建Agent...")
    agent = create_agent_instance()
    print("[INFO] Agent创建成功！")
    
    # 测试调用
    config = {"configurable": {"thread_id": "test_1"}}
    result = agent.invoke(
        {"messages": [{"role": "user", "content": "查看集群状态"}]},
        config=config
    )
    
    print("\n[结果]")
    print(result['messages'][-1].content)

