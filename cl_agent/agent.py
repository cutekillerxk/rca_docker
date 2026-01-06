#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HDFS集群监控Agent（LangChain 1.0.7 + vLLM）
整合所有功能：日志读取、监控采集、Agent创建、文档导出
"""

import os
import json
import re
import logging
import subprocess
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
import paramiko
from dotenv import load_dotenv

load_dotenv()

# LangChain 导入
from langchain.agents import create_agent
from langchain.tools import tool
from langchain_openai import ChatOpenAI

from langchain_tavily import TavilySearch
website_search = TavilySearch(max_results=2, tavily_api_key=os.getenv("TAVILY_API_KEY"))

from langchain_community.tools.shell.tool import ShellTool

shell = ShellTool()
# ==================== 配置常量 ====================

# JMX API 地址
# Docker环境：本机Docker容器端口映射
NAMENODE = "http://localhost:9870/jmx"
DATANODES = [
    "http://127.0.0.1:9864/jmx",  # datanode1
    "http://127.0.0.1:9865/jmx",  # datanode2
    "http://127.0.0.1:9866/jmx",  # datanode-namenode
]

# 日志文件配置
# Docker环境：从容器内日志文件读取
# 日志路径：/usr/local/hadoop/logs/ (HADOOP_HOME=/usr/local/hadoop)
# 注意：namenode容器中运行了3个服务：namenode, datanode, secondarynamenode
# 所以需要读取5个日志文件：
# 1. namenode容器：hadoop-hadoop-namenode-namenode.log
# 2. namenode容器：hadoop-hadoop-datanode-namenode.log
# 3. namenode容器：hadoop-hadoop-secondarynamenode-namenode.log
# 4. datanode1容器：hadoop-hadoop-datanode-datanode1.log
# 5. datanode2容器：hadoop-hadoop-datanode-datanode2.log
LOG_FILES_CONFIG = [
    {
        "name": "namenode",
        "display_name": "NameNode",
        "type": "docker",
        "container": "namenode",
        "log_path": "/usr/local/hadoop/logs",
        "node_pattern": "namenode",  # 匹配 hadoop-hadoop-namenode-namenode.log
        "ip": "127.0.0.1"
    },
    {
        "name": "datanode-namenode",
        "display_name": "DataNode (NameNode容器)",
        "type": "docker",
        "container": "namenode",
        "log_path": "/usr/local/hadoop/logs",
        "node_pattern": "datanode",  # 匹配 hadoop-hadoop-datanode-namenode.log
        "ip": "127.0.0.1"
    },
    {
        "name": "secondarynamenode",
        "display_name": "SecondaryNameNode",
        "type": "docker",
        "container": "namenode",
        "log_path": "/usr/local/hadoop/logs",
        "node_pattern": "secondarynamenode",  # 匹配 hadoop-hadoop-secondarynamenode-namenode.log
        "ip": "127.0.0.1"
    },
    {
        "name": "datanode1",
        "display_name": "DataNode1",
        "type": "docker",
        "container": "datanode1",
        "log_path": "/usr/local/hadoop/logs",
        "node_pattern": "datanode",  # 匹配 hadoop-hadoop-datanode-datanode1.log
        "ip": "127.0.0.1"
    },
    {
        "name": "datanode2",
        "display_name": "DataNode2",
        "type": "docker",
        "container": "datanode2",
        "log_path": "/usr/local/hadoop/logs",
        "node_pattern": "datanode",  # 匹配 hadoop-hadoop-datanode-datanode2.log
        "ip": "127.0.0.1"
    }
]

# Docker 容器配置（替代原来的SSH配置）
# 日志路径已确认：/usr/local/hadoop/logs/ (HADOOP_HOME=/usr/local/hadoop)
# 注意：namenode容器中运行了3个服务，所以有3个日志文件
DOCKER_CONFIG = {
    "namenode": {
        "container": "namenode",
        "log_path": "/usr/local/hadoop/logs",
    },
    "datanode-namenode": {
        "container": "namenode",
        "log_path": "/usr/local/hadoop/logs",
    },
    "secondarynamenode": {
        "container": "namenode",
        "log_path": "/usr/local/hadoop/logs",
    },
    "datanode1": {
        "container": "datanode1",
        "log_path": "/usr/local/hadoop/logs",
    },
    "datanode2": {
        "container": "datanode2",
        "log_path": "/usr/local/hadoop/logs",
    }
}

# 保留SSH配置（已废弃，但保留以防回滚需要）
SSH_CONFIG = {
    # 已迁移到Docker环境，此配置不再使用
}

# 路径配置
LOG_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "logs")
STATE_FILE = os.path.join(LOG_DIR, "log_reader_state.json")
os.makedirs(LOG_DIR, exist_ok=True)

# 日志过滤配置
FILTER_INFO_LOGS = True      # 是否过滤INFO级别日志
FILTER_CLASSPATH_LOGS = True  # 是否过滤classpath行

# 日志读取配置
DEFAULT_MAX_LINES = 200  # 默认读取日志行数

# vLLM 配置
VLLM_BASE_URL = "http://10.157.197.76:8001/v1"
VLLM_MODEL_PATH = "/media/hnu/LLM/hnu/LLM/Qwen3-8B"

# 第三方 API 配置（从环境变量读取）
THIRD_PARTY_API_BASE_URL = os.getenv("API_BASE_URL", "")  # 第三方 API base_url
THIRD_PARTY_API_KEY = os.getenv("API_KEY", "")  # 第三方 API key（DeepSeek 和 GPT 共用）

# 全局变量
ssh_readers = {}  # 保留用于兼容性
docker_readers = {}  # Docker日志读取器

# ==================== Docker 日志读取器类 ====================

class DockerLogReader:
    """通过 Docker exec 读取容器日志文件"""
    
    def __init__(self, container: str, log_path: str):
        self.container = container
        self.log_path = log_path
        self._connected = True  # Docker exec不需要持久连接
    
    def list_log_files(self, node_pattern: Optional[str] = None) -> List[str]:
        """列出容器中的日志文件"""
        try:
            # 首先检查容器是否运行（跨平台方式：使用Python直接检查）
            check_result = subprocess.run(
                ['docker', 'ps', '--format', '{{.Names}}'],
                shell=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                universal_newlines=True,
                timeout=5
            )
            
            if check_result.returncode != 0 or self.container not in check_result.stdout:
                logging.warning(f"容器 {self.container} 未运行，无法列出日志文件")
                return []
            
            # 使用docker exec执行ls命令（兼容Windows PowerShell）
            # 先尝试列出所有文件，包括.log和.audit文件
            result = subprocess.run(
                f'docker exec {self.container} sh -c "ls -1 {self.log_path} 2>&1"',
                shell=True,
                capture_output=True,
                text=True,
                timeout=10
            )
            
            if result.returncode != 0:
                logging.warning(f"无法列出容器 {self.container} 的日志文件: {result.stderr}")
                return []
            
            files = result.stdout.strip().split('\n')
            files = [f for f in files if f.strip()]  # 移除空行
            # 包含.log和.audit文件（Hadoop的日志文件）
            log_files = [f for f in files if f.endswith(".log") or f.endswith(".audit") or f.endswith(".out")]
            
            if node_pattern:
                log_files = [f for f in log_files if node_pattern.lower() in f.lower()]
            
            return log_files
        except Exception as e:
            logging.error(f"列出日志文件失败: {e}")
            return []
    
    def read_log_file(self, file_path: str, start_pos: int = 0, 
                     max_lines: Optional[int] = None) -> Tuple[str, int]:
        """从容器文件读取日志"""
        try:
            # 如果文件路径不是绝对路径，拼接日志目录
            if not os.path.isabs(file_path):
                file_path = os.path.join(self.log_path, file_path)
            
            # 确保路径使用正斜杠（容器内路径）
            file_path = file_path.replace('\\', '/')
            
            # 使用docker exec执行tail命令读取日志（兼容Windows PowerShell）
            # 使用字节位置确保完全准确
            escaped_path = file_path.replace("'", "'\"'\"'")
            
            if max_lines:
                # 读取指定行数（从文件末尾或指定位置）
                if start_pos == 0:
                    # 从文件末尾读取
                    result = subprocess.run(
                        f'docker exec {self.container} sh -c "tail -n {max_lines} \'{escaped_path}\' 2>&1"',
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=10
                    )
                else:
                    # 从指定字节位置读取，然后在Python中限制行数（完全准确）
                    result = subprocess.run(
                        f'docker exec {self.container} sh -c "tail -c +{start_pos} \'{escaped_path}\' 2>&1"',
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=10
                    )
            else:
                # 读取全部（从指定位置）
                if start_pos == 0:
                    result = subprocess.run(
                        f'docker exec {self.container} sh -c "cat \'{escaped_path}\' 2>&1"',
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=10
                    )
                else:
                    result = subprocess.run(
                        f'docker exec {self.container} sh -c "tail -c +{start_pos} \'{escaped_path}\' 2>&1"',
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=10
                    )
            
            if result.returncode != 0:
                error_msg = result.stderr if result.stderr else result.stdout
                logging.error(f"读取日志文件失败 (容器: {self.container}, 文件: {file_path}, 返回码: {result.returncode}): {error_msg}")
                return "", start_pos
            
            content = result.stdout
            
            # 如果指定了max_lines且从指定位置读取，需要在Python中限制行数
            if max_lines and start_pos != 0:
                lines = content.split('\n')
                content = '\n'.join(lines[:max_lines])
            
            # 计算新位置：使用实际读取内容的字节数（完全准确）
            new_pos = start_pos + len(content.encode('utf-8'))
            
            return content, new_pos
            
        except Exception as e:
            logging.error(f"读取日志文件失败 (容器: {self.container}, 文件: {file_path}): {e}")
            import traceback
            logging.debug(traceback.format_exc())
            return "", start_pos
    
    def get_file_mtime(self, file_path: str) -> Optional[float]:
        """获取容器内文件的修改时间"""
        try:
            if not os.path.isabs(file_path):
                file_path = os.path.join(self.log_path, file_path)
            
            result = subprocess.run(
                f'docker exec {self.container} sh -c "stat -c %Y {file_path} 2>&1"',
                shell=True,
                capture_output=True,
                text=True,
                timeout=10
            )
            
            if result.returncode == 0:
                try:
                    return float(result.stdout.strip())
                except ValueError:
                    return None
            return None
        except Exception as e:
            logging.error(f"获取文件修改时间失败: {e}")
            return None
    
    def check_file_exists(self, file_path: str) -> bool:
        """检查容器内文件是否存在"""
        try:
            if not os.path.isabs(file_path):
                file_path = os.path.join(self.log_path, file_path)
            
            result = subprocess.run(
                f'docker exec {self.container} sh -c "test -f {file_path} 2>&1"',
                shell=True,
                capture_output=True,
                text=True,
                timeout=10
            )
            return result.returncode == 0
        except Exception as e:
            logging.error(f"检查文件是否存在失败: {e}")
            return False
    
    def read_docker_logs(self, max_lines: Optional[int] = None, tail: bool = True) -> str:
        """
        读取Docker容器的标准输出日志（docker logs）
        
        Args:
            max_lines: 最大读取行数
            tail: 是否从末尾读取（True）还是从头读取（False）
        
        Returns:
            日志内容
        """
        try:
            # 首先检查容器是否存在（即使未运行，docker logs也可以读取历史日志）
            # 使用跨平台方式：直接使用Python检查
            check_result = subprocess.run(
                ['docker', 'ps', '-a', '--format', '{{.Names}}'],
                shell=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                universal_newlines=True,
                timeout=5
            )
            
            if check_result.returncode != 0 or self.container not in check_result.stdout:
                logging.warning(f"容器 {self.container} 不存在，无法读取日志")
                return ""
            if max_lines:
                if tail:
                    # 从末尾读取指定行数
                    result = subprocess.run(
                        f'docker logs {self.container} --tail {max_lines} 2>&1',
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=10
                    )
                else:
                    # 从头读取指定行数（Windows兼容：先读取全部，再截取前N行）
                    result = subprocess.run(
                        f'docker logs {self.container} 2>&1',
                        shell=True,
                        capture_output=True,
                        text=True,
                        timeout=10
                    )
                    if result.returncode == 0:
                        # 在Python中截取前max_lines行
                        lines = result.stdout.splitlines()
                        result.stdout = '\n'.join(lines[:max_lines])
            else:
                # 读取全部日志
                result = subprocess.run(
                    f'docker logs {self.container} 2>&1',
                    shell=True,
                    capture_output=True,
                    text=True,
                    timeout=10
                )
            
            if result.returncode == 0:
                return result.stdout
            else:
                logging.error(f"读取Docker日志失败: {result.stderr}")
                return ""
        except Exception as e:
            logging.error(f"读取Docker日志失败: {e}")
            return ""

# ==================== SSH 日志读取器类 ====================

class SSHLogReader:
    """通过 SSH 读取远程日志文件"""
    
    def __init__(self, host: str, user: str, log_path: str, port: int = 22, 
                 key_file: Optional[str] = None, password: Optional[str] = None):
        self.host = host
        self.user = user
        self.port = port
        self.log_path = log_path
        self.key_file = key_file
        self.password = password
        self.client = None
        self.sftp = None
        self._connected = False
    
    def connect(self) -> bool:
        """建立 SSH 连接"""
        if self._connected and self.client and self.client.get_transport() and self.client.get_transport().is_active():
            return True
        
        try:
            self.client = paramiko.SSHClient()
            self.client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
            
            if self.key_file:
                key_path = os.path.expanduser(self.key_file)
                if os.path.exists(key_path):
                    if not os.access(key_path, os.R_OK):
                        raise PermissionError(f"无权限读取 SSH 密钥文件: {key_path}")
                    
                    try:
                        private_key = paramiko.RSAKey.from_private_key_file(key_path)
                        self.client.connect(
                            hostname=self.host,
                            port=self.port,
                            username=self.user,
                            pkey=private_key,
                            timeout=10,
                            look_for_keys=False
                        )
                    except Exception as e:
                        logging.warning(f"使用指定密钥文件失败: {e}，尝试使用默认密钥")
                        self.client.connect(
                            hostname=self.host,
                            port=self.port,
                            username=self.user,
                            timeout=10
                        )
                else:
                    self.client.connect(
                        hostname=self.host,
                        port=self.port,
                        username=self.user,
                        timeout=10
                    )
            elif self.password:
                self.client.connect(
                    hostname=self.host,
                    port=self.port,
                    username=self.user,
                    password=self.password,
                    timeout=10
                )
            else:
                self.client.connect(
                    hostname=self.host,
                    port=self.port,
                    username=self.user,
                    timeout=10
                )
            
            self.sftp = self.client.open_sftp()
            self._connected = True
            return True
            
        except Exception as e:
            logging.error(f"Failed to connect to {self.host}: {e}")
            return False
    
    def disconnect(self):
        """关闭 SSH 连接"""
        try:
            if self.sftp:
                self.sftp.close()
                self.sftp = None
            if self.client:
                self.client.close()
                self.client = None
            self._connected = False
        except Exception as e:
            logging.warning(f"Error closing SSH connection: {e}")
    
    def execute_command(self, command: str, timeout: int = 60) -> Tuple[str, str, int]:
        """
        通过SSH执行远程命令
        
        Args:
            command: 要执行的命令
            timeout: 超时时间（秒）
        
        Returns:
            (stdout, stderr, returncode) 元组
        """
        if not self._connected:
            if not self.connect():
                return "", f"无法连接到 {self.host}", -1
        
        try:
            # 设置Hadoop环境变量
            hadoop_home = "/media/hnu/dependency/hadoop/module/hadoop-3.1.3"
            env_command = f"export HADOOP_HOME={hadoop_home} && export PATH={hadoop_home}/bin:{hadoop_home}/sbin:$PATH && {command}"
            
            stdin, stdout, stderr = self.client.exec_command(env_command, timeout=timeout)
            
            # 读取输出
            stdout_text = stdout.read().decode('utf-8', errors='ignore')
            stderr_text = stderr.read().decode('utf-8', errors='ignore')
            returncode = stdout.channel.recv_exit_status()
            
            return stdout_text, stderr_text, returncode
            
        except Exception as e:
            error_msg = f"执行远程命令失败: {str(e)}"
            logging.error(f"[{self.host}] {error_msg}")
            return "", error_msg, -1
    
    def list_log_files(self, node_pattern: Optional[str] = None) -> List[str]:
        """列出远程目录中的日志文件"""
        if not self._connected:
            if not self.connect():
                return []
        
        try:
            files = self.sftp.listdir(self.log_path)
            log_files = [f for f in files if f.endswith(".log")]
            
            if node_pattern:
                log_files = [f for f in log_files if node_pattern.lower() in f.lower()]
            
            return log_files
        except Exception as e:
            logging.error(f"Failed to list files in {self.log_path}: {e}")
            return []
    
    def read_log_file(self, file_path: str, start_pos: int = 0, 
                     max_lines: Optional[int] = None) -> Tuple[str, int]:
        """从远程文件读取日志"""
        if not self._connected:
            if not self.connect():
                return "", start_pos
        
        try:
            if not os.path.isabs(file_path):
                file_path = os.path.join(self.log_path, file_path)
            
            try:
                file_stat = self.sftp.stat(file_path)
                file_size = file_stat.st_size
                if start_pos > file_size:
                    logging.info(f"日志文件可能被轮转，重置读取位置: {file_path}")
                    start_pos = 0
            except Exception:
                file_size = None
            
            remote_file = self.sftp.open(file_path, 'r')
            remote_file.seek(start_pos)
            
            if max_lines:
                lines = []
                initial_pos = remote_file.tell()
                
                for i in range(max_lines):
                    line = remote_file.readline()
                    if not line:
                        break
                    if isinstance(line, bytes):
                        line_str = line.decode('utf-8', errors='ignore')
                    else:
                        line_str = line
                    
                    # 过滤日志
                    if not should_filter_log_line(line_str, FILTER_INFO_LOGS, FILTER_CLASSPATH_LOGS):
                        lines.append(line_str)
                
                content = ''.join(lines)
                current_pos = remote_file.tell()
                
                if current_pos == initial_pos and not content:
                    if file_size is not None:
                        current_pos = file_size
            else:
                content_parts = []
                max_default_lines = 500
                lines_read = 0
                while lines_read < max_default_lines:
                    line = remote_file.readline()
                    if not line:
                        break
                    if isinstance(line, bytes):
                        line_str = line.decode('utf-8', errors='ignore')
                    else:
                        line_str = line
                    
                    # 过滤日志
                    if not should_filter_log_line(line_str, FILTER_INFO_LOGS, FILTER_CLASSPATH_LOGS):
                        content_parts.append(line_str)
                    lines_read += 1
                
                content = ''.join(content_parts)
                current_pos = remote_file.tell()
            
            remote_file.close()
            return content, current_pos
            
        except Exception as e:
            logging.error(f"Failed to read file {file_path}: {e}")
            return "", start_pos
    
    def get_file_mtime(self, file_path: str) -> Optional[float]:
        """获取远程文件的修改时间"""
        if not self._connected:
            if not self.connect():
                return None
        
        try:
            if not os.path.isabs(file_path):
                file_path = os.path.join(self.log_path, file_path)
            
            file_stat = self.sftp.stat(file_path)
            return file_stat.st_mtime
        except Exception as e:
            logging.error(f"Failed to get mtime for {file_path}: {e}")
            return None
    
    def check_file_exists(self, file_path: str) -> bool:
        """检查远程文件是否存在"""
        if not self._connected:
            if not self.connect():
                return False
        
        try:
            if not os.path.isabs(file_path):
                file_path = os.path.join(self.log_path, file_path)
            self.sftp.stat(file_path)
            return True
        except FileNotFoundError:
            return False
        except Exception as e:
            logging.error(f"Error checking file {file_path}: {e}")
            return False

# ==================== 日志读取功能 ====================

def should_filter_log_line(line: str, filter_info: bool = True, filter_classpath: bool = True) -> bool:
    """
    判断是否应该过滤该日志行
    
    Args:
        line: 日志行内容
        filter_info: 是否过滤INFO级别日志
        filter_classpath: 是否过滤classpath行
    
    Returns:
        True表示应该过滤（跳过），False表示保留
    """
    if not line.strip():
        return False  # 空行保留
    
    # 过滤classpath行
    if filter_classpath:
        if re.search(r'STARTUP_MSG:\s+classpath\s*=', line, re.IGNORECASE):
            return True
    
    # 过滤INFO级别日志
    if filter_info:
        # 匹配标准格式：时间戳 + INFO
        if re.search(r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}[,\s]+\d+\s+INFO\s+', line):
            return True
    
    return False  # 保留该行


def read_latest_logs(path: str, last_pos: int, node_pattern: Optional[str] = None, max_lines: int = DEFAULT_MAX_LINES) -> Tuple[List[str], int]:
    """滚动读取最新日志（本地）"""
    print(f"读取最新日志: {path}")
    try:
        if not os.path.exists(path) or not os.path.isdir(path):
            return [], last_pos
        
        if not os.access(path, os.R_OK):
            logging.warning(f"无读取权限: {path}")
            return [], last_pos
        
        try:
            all_log_files = [f for f in os.listdir(path) if f.endswith(".log")]
        except PermissionError:
            logging.warning(f"无权限列出目录内容: {path}")
            return [], last_pos
        
        if node_pattern:
            log_files = [f for f in all_log_files if node_pattern.lower() in f.lower()]
        else:
            log_files = all_log_files
        
        if not log_files:
            return [], last_pos
        
        log_files_with_time = []
        for f in log_files:
            try:
                file_path = os.path.join(path, f)
                mtime = os.path.getmtime(file_path)
                log_files_with_time.append((f, mtime))
            except (OSError, PermissionError):
                continue
        
        if not log_files_with_time:
            return [], last_pos
        
        log_files_with_time.sort(key=lambda x: x[1], reverse=True)
        latest_file = os.path.join(path, log_files_with_time[0][0])
        
        if not os.access(latest_file, os.R_OK):
            logging.warning(f"无读取权限: {latest_file}")
            return [], last_pos
        
        current_size = os.path.getsize(latest_file)
        if last_pos > current_size:
            logging.info(f"日志文件可能被轮转，重置读取位置: {latest_file}")
            last_pos = 0
        
        with open(latest_file, "r", encoding='utf-8', errors='ignore') as f:
            f.seek(last_pos)
            lines = [] 
            for i in range(max_lines):
                line = f.readline()
                if not line:
                    break
                # 过滤日志
                if not should_filter_log_line(line, FILTER_INFO_LOGS, FILTER_CLASSPATH_LOGS):
                    lines.append(line)
            new_pos = f.tell()
        
        return lines, new_pos
        
    except Exception as e:
        logging.warning(f"Failed to read logs from {path}: {e}")
        return [], last_pos


def read_latest_logs_ssh(ssh_reader: SSHLogReader, last_pos: int, 
                        node_pattern: Optional[str] = None, max_lines: int = DEFAULT_MAX_LINES,
                        last_file: Optional[str] = None) -> Tuple[List[str], int, Optional[str]]:
    """通过 SSH 读取远程最新日志"""
    print(f"通过 SSH 读取远程最新日志: {ssh_reader.host}")
    try:
        if not ssh_reader.connect():
            return [], last_pos, last_file
        
        log_files = ssh_reader.list_log_files(node_pattern)
        if not log_files:
            return [], last_pos, last_file
        
        log_files_with_time = []
        for f in log_files:
            mtime = ssh_reader.get_file_mtime(f)
            if mtime:
                log_files_with_time.append((f, mtime))
        
        if not log_files_with_time:
            return [], last_pos, last_file
        
        log_files_with_time.sort(key=lambda x: x[1], reverse=True)
        latest_file = log_files_with_time[0][0]
        
        if last_file and last_file != latest_file:
            logging.info(f"检测到日志文件切换: {last_file} -> {latest_file}，重置读取位置")
            last_pos = 0
        
        content, new_pos = ssh_reader.read_log_file(latest_file, last_pos, max_lines=max_lines)
        lines = content.splitlines(keepends=True) if content else []
        
        return lines, new_pos, latest_file
        
    except Exception as e:
        logging.error(f"Failed to read logs via SSH from {ssh_reader.host}: {e}")
        return [], last_pos, last_file


def read_latest_logs_docker(docker_reader: DockerLogReader, last_pos: int,
                            node_pattern: Optional[str] = None, max_lines: int = DEFAULT_MAX_LINES,
                            last_file: Optional[str] = None) -> Tuple[List[str], int, Optional[str]]:
    """通过 Docker exec 读取容器最新日志（强制从日志文件读取，不回退到docker_logs）"""
    print(f"通过 Docker 读取容器日志: {docker_reader.container}")
    try:
        # 强制从日志文件读取，不回退到docker_logs
        log_files = docker_reader.list_log_files(node_pattern)
        
        if not log_files:
            # 如果没有找到日志文件，返回错误信息
            error_msg = f"未找到日志文件（容器: {docker_reader.container}, 路径: {docker_reader.log_path}, 模式: {node_pattern}）"
            logging.warning(error_msg)
            return [], last_pos, last_file
        
        # 有日志文件，尝试读取文件
        log_files_with_time = []
        for f in log_files:
            mtime = docker_reader.get_file_mtime(f)
            if mtime:
                log_files_with_time.append((f, mtime))
        
        if not log_files_with_time:
            # 如果无法获取文件时间，使用第一个文件
            logging.warning(f"无法获取日志文件修改时间，使用第一个文件: {log_files[0]}")
            latest_file = log_files[0]
        else:
            log_files_with_time.sort(key=lambda x: x[1], reverse=True)
            latest_file = log_files_with_time[0][0]
        
        if last_file and last_file != latest_file:
            logging.info(f"检测到日志文件切换: {last_file} -> {latest_file}，重置读取位置")
            last_pos = 0
        
        # 读取日志文件内容
        content, new_pos = docker_reader.read_log_file(latest_file, last_pos, max_lines=max_lines)
        
        if not content:
            # 文件为空，返回空列表
            logging.info(f"日志文件 {latest_file} 为空")
            return [], new_pos, latest_file
        
        lines = content.splitlines(keepends=True)
        
        # 过滤日志（去除INFO级别和classpath行）
        filtered_lines = []
        for line in lines:
            if not should_filter_log_line(line, FILTER_INFO_LOGS, FILTER_CLASSPATH_LOGS):
                filtered_lines.append(line)
        
        # 如果过滤后有内容，返回过滤后的内容；否则返回原始内容
        if filtered_lines:
            return filtered_lines, new_pos, latest_file
        else:
            # 过滤后没有内容，返回原始内容（不过滤）
            logging.info(f"日志文件 {latest_file} 过滤后无内容，返回原始内容")
            return lines, new_pos, latest_file
        
    except Exception as e:
        error_msg = f"读取容器日志文件失败（容器: {docker_reader.container}, 路径: {docker_reader.log_path}）: {str(e)}"
        logging.error(error_msg)
        # 出错时返回空列表，不回退到docker_logs
        return [], last_pos, last_file


def init_ssh_readers():
    """初始化 SSH 读取器（已废弃，保留用于兼容性）"""
    print("初始化 SSH 读取器（已废弃）")
    global ssh_readers
    
    for log_config in LOG_FILES_CONFIG:
        if log_config["type"] == "ssh":
            host = log_config["host"]
            if host not in ssh_readers:
                key_file = None
                for key, config in SSH_CONFIG.items():
                    if config.get("host") == host:
                        key_file = config.get("key_file")
                        break
                if not key_file:
                    key_file = "/home/hadoop/.ssh/id_rsa"
                
                ssh_readers[host] = SSHLogReader(
                    host=host,
                    user=log_config["user"],
                    log_path=log_config["log_path"],
                    key_file=key_file
                )


def init_docker_readers():
    """初始化 Docker 读取器"""
    print("初始化 Docker 读取器")
    global docker_readers
    
    for log_config in LOG_FILES_CONFIG:
        if log_config["type"] == "docker":
            container = log_config.get("container")
            if container and container not in docker_readers:
                docker_readers[container] = DockerLogReader(
                    container=container,
                    log_path=log_config["log_path"]
                )


def load_log_reader_state(num_log_files: int) -> Tuple[List[int], List[Optional[str]]]:
    """从文件加载日志读取器的状态"""
    print("从文件加载日志读取器的状态")
    try:
        if os.path.exists(STATE_FILE):
            with open(STATE_FILE, 'r', encoding='utf-8') as f:
                state = json.load(f)
                last_positions = state.get('last_positions', [0] * num_log_files)
                last_files = state.get('last_files', [None] * num_log_files)
                
                if len(last_positions) != num_log_files:
                    last_positions = [0] * num_log_files
                if len(last_files) != num_log_files:
                    last_files = [None] * num_log_files
                
                return last_positions, last_files
        else:
            return [0] * num_log_files, [None] * num_log_files
    except Exception as e:
        logging.warning(f"加载状态文件失败: {e}，使用默认值")
        return [0] * num_log_files, [None] * num_log_files


def save_log_reader_state(last_positions: List[int], last_files: List[Optional[str]]):
    """保存日志读取器的状态到文件"""
    print("保存日志读取器的状态到文件")
    try:
        os.makedirs(LOG_DIR, exist_ok=True)
        
        state = {
            'last_positions': last_positions,
            'last_files': last_files,
            'last_update': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
        
        with open(STATE_FILE, 'w', encoding='utf-8') as f:
            json.dump(state, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logging.warning(f"保存状态文件失败: {e}")


def read_all_cluster_logs(max_lines: int = DEFAULT_MAX_LINES, 
                          last_positions: Optional[List[int]] = None,
                          last_files: Optional[List[Optional[str]]] = None) -> Tuple[Dict[str, str], List[int], List[Optional[str]]]:
    """读取所有5个节点的日志"""
    print("读取所有5个节点的日志")
    logs = {}
    num_log_files = len(LOG_FILES_CONFIG)
    
    if last_positions is None:
        last_positions = [0] * num_log_files
    if last_files is None:
        last_files = [None] * num_log_files
    
    if len(last_positions) != num_log_files:
        last_positions = [0] * num_log_files
    if len(last_files) != num_log_files:
        last_files = [None] * num_log_files
    
    init_ssh_readers()  # 保留用于兼容性
    init_docker_readers()  # 初始化Docker读取器
    
    new_positions = []
    new_files = []
    
    for i, log_config in enumerate(LOG_FILES_CONFIG):
        node_name = log_config["display_name"]
        log_path = log_config["log_path"]
        node_pattern = log_config.get("node_pattern")
        
        try:
            if log_config["type"] == "local":
                lines, new_pos = read_latest_logs(
                    log_path, 
                    last_pos=last_positions[i],
                    node_pattern=node_pattern,
                    max_lines=max_lines
                )
                log_content = "".join(lines)
                new_positions.append(new_pos)
                new_files.append(None)
            elif log_config["type"] == "docker":
                container = log_config.get("container")
                docker_reader = docker_readers.get(container) if container else None
                
                if docker_reader:
                    lines, new_pos, current_file = read_latest_logs_docker(
                        docker_reader,
                        last_pos=last_positions[i],
                        node_pattern=node_pattern,
                        max_lines=max_lines,
                        last_file=last_files[i]
                    )
                    log_content = "".join(lines)
                    new_positions.append(new_pos)
                    new_files.append(current_file)
                else:
                    log_content = f"无法连接到容器 {container}"
                    new_positions.append(last_positions[i])
                    new_files.append(last_files[i])
            else:  # ssh类型（已废弃，保留用于兼容性）
                host = log_config.get("host")
                ssh_reader = ssh_readers.get(host) if host else None
                
                if ssh_reader:
                    lines, new_pos, current_file = read_latest_logs_ssh(
                        ssh_reader,
                        last_pos=last_positions[i],
                        node_pattern=node_pattern,
                        max_lines=max_lines,
                        last_file=last_files[i]
                    )
                    log_content = "".join(lines)
                    new_positions.append(new_pos)
                    new_files.append(current_file)
                else:
                    log_content = f"无法连接到 {host}"
                    new_positions.append(last_positions[i])
                    new_files.append(last_files[i])
            
            logs[node_name] = log_content if log_content else "无新日志"
            
        except Exception as e:
            logs[node_name] = f"读取日志失败: {str(e)}"
            new_positions.append(last_positions[i])
            new_files.append(last_files[i])
    
    return logs, new_positions, new_files


def get_node_log_by_name(node_name: str) -> str:
    """根据节点名称获取日志"""
    print(f"根据节点名称获取日志: {node_name}")
    node_name_lower = node_name.lower()
    target_config = None
    
    for config in LOG_FILES_CONFIG:
        display_name = config.get("display_name", "").lower()
        if (node_name_lower in display_name or 
            display_name in node_name_lower or
            (node_name_lower == "namenode" and "namenode" in display_name) or
            (node_name_lower == "secondarynamenode" and "secondarynamenode" in display_name)):
            target_config = config
            break
    
    if not target_config:
        # 更新支持的节点列表
        supported_nodes = [config["display_name"] for config in LOG_FILES_CONFIG]
        return f"未找到节点: {node_name}。支持的节点: {', '.join(supported_nodes)}"
    
    # 加载状态
    num_log_files = len(LOG_FILES_CONFIG)
    last_positions, last_files = load_log_reader_state(num_log_files)
    
    # 找到目标节点的索引
    target_idx = None
    for i, config in enumerate(LOG_FILES_CONFIG):
        if config == target_config:
            target_idx = i
            break
    
    if target_idx is None:
        return f"未找到节点配置: {node_name}"
    
    # 读取该节点的日志
    log_path = target_config["log_path"]
    node_pattern = target_config.get("node_pattern")
    
    try:
        if target_config["type"] == "local":
            lines, _ = read_latest_logs(
                log_path,
                last_pos=last_positions[target_idx],
                node_pattern=node_pattern,
                max_lines=DEFAULT_MAX_LINES
            )
            log_content = "".join(lines)
        elif target_config["type"] == "docker":
            container = target_config.get("container")
            init_docker_readers()
            docker_reader = docker_readers.get(container) if container else None
            
            if docker_reader:
                lines, _, _ = read_latest_logs_docker(
                    docker_reader,
                    last_pos=last_positions[target_idx],
                    node_pattern=node_pattern,
                    max_lines=DEFAULT_MAX_LINES,
                    last_file=last_files[target_idx]
                )
                log_content = "".join(lines)
            else:
                log_content = f"无法连接到容器 {container}"
        else:  # ssh类型（已废弃，保留用于兼容性）
            host = target_config.get("host")
            init_ssh_readers()
            ssh_reader = ssh_readers.get(host) if host else None
            
            if ssh_reader:
                lines, _, _ = read_latest_logs_ssh(
                    ssh_reader,
                    last_pos=last_positions[target_idx],
                    node_pattern=node_pattern,
                    max_lines=DEFAULT_MAX_LINES,
                    last_file=last_files[target_idx]
                )
                log_content = "".join(lines)
            else:
                log_content = f"无法连接到 {host}"
        
        return log_content if log_content else "无新日志"
    except Exception as e:
        return f"读取日志失败: {str(e)}"

# ==================== LLM 配置 ====================

def create_llm(model_name: str = "qwen-8b") -> ChatOpenAI:
    """
    根据模型名称创建LLM实例
    
    Args:
        model_name: 模型名称，可选值：
            - "qwen-8b": Qwen-8B (vLLM本地部署)
            - "gpt-4o": GPT-4o (第三方API)
            - "deepseek-r1": DeepSeek-R1 (第三方API)
    
    Returns:
        ChatOpenAI 实例
    """
    # 模型配置字典
    model_configs = {
        "qwen-8b": {
            "base_url": VLLM_BASE_URL,
            "api_key": "not-needed",
            "model": VLLM_MODEL_PATH,
            "timeout": 120,
            "max_tokens": 4096,
        },
        "gpt-4o": {
            "base_url": THIRD_PARTY_API_BASE_URL,
            "api_key": THIRD_PARTY_API_KEY,
            "model": "gpt-4o",  # 实际模型名根据第三方API调整
            "timeout": 60,
            "max_tokens": 4096,
        },
        "deepseek-r1": {
            "base_url": THIRD_PARTY_API_BASE_URL,
            "api_key": THIRD_PARTY_API_KEY,
            "model": "DeepSeek-V3.2",  # 实际模型名根据第三方API调整
            "timeout": 120,
            "max_tokens": 4096,
        }
    }
    
    # 获取模型配置，如果不存在则使用默认配置（qwen-8b）
    config = model_configs.get(model_name, model_configs["qwen-8b"])
    
    # 检查第三方API配置（gpt-4o 和 deepseek-r1 需要）
    if model_name in ["gpt-4o", "deepseek-r1"]:
        if not config["base_url"] or not config["api_key"]:
            raise ValueError(
                f"模型 {model_name} 需要配置 API_BASE_URL 和 API_KEY 环境变量。\n"
                f"请在 .env 文件中设置：\n"
                f"  API_BASE_URL=你的第三方API地址\n"
                f"  API_KEY=你的API密钥"
            )
    
    # 调试信息：打印模型配置
    print(f"[DEBUG] 创建LLM实例 - 模型: {model_name}")
    print(f"[DEBUG]   - base_url: {config['base_url']}")
    print(f"[DEBUG]   - model: {config['model']}")
    # 安全显示 API key（只显示后4位）
    if config['api_key'] == "not-needed":
        api_key_display = "not-needed (vLLM本地部署)"
    elif config['api_key']:
        api_key_display = f"已设置 (***{config['api_key'][-4:]})" if len(config['api_key']) > 4 else "已设置 (***)"
    else:
        api_key_display = "未设置"
    print(f"[DEBUG]   - api_key: {api_key_display}")
    print(f"[DEBUG]   - timeout: {config['timeout']}s")
    print(f"[DEBUG]   - max_tokens: {config['max_tokens']}")
    
    llm = ChatOpenAI(
        base_url=config["base_url"],
        api_key=config["api_key"],
        model=config["model"],
        temperature=0,
        max_tokens=config["max_tokens"],
        timeout=config["timeout"],
        max_retries=2,
    )
    
    # 调试信息：确认LLM实例创建成功
    print(f"[DEBUG] ✅ LLM实例创建成功 - 模型: {model_name} (实际模型名: {config['model']})")
    
    return llm

# ==================== 工具函数定义 ====================

@tool("get_cluster_logs", description="获取集群所有节点的最新日志内容，用来分析集群的整体状态")
def get_cluster_logs() -> str:
    """
    获取所有5个节点的日志内容。
    返回格式化的日志文本，由Agent逐个分析每个节点。
    
    Returns:
        格式化的日志文本，包含思考检查点
    """
    print("调用get_cluster_logs工具")
    try:
        num_log_files = len(LOG_FILES_CONFIG)
        last_positions, last_files = load_log_reader_state(num_log_files)
        
        all_logs, new_positions, new_files = read_all_cluster_logs(
            max_lines=DEFAULT_MAX_LINES,
            last_positions=last_positions,
            last_files=last_files
        )
        
        save_log_reader_state(new_positions, new_files)
        
        # 构建带思考检查点的返回内容
        result = []
        result.append("[集群日志分析任务]")
        result.append(f"共发现 {len(all_logs)} 个节点需要分析。")
        result.append("请逐个分析每个节点的日志，每个节点分析一次。")
        result.append("分析完一个节点后，记录分析结果，然后继续下一个节点。")
        result.append("")
        
        node_list = list(all_logs.items())
        for idx, (node_name, log_content) in enumerate(node_list, 1):
            result.append(f"=== {node_name} ===")
            result.append(log_content)
            result.append("")
            
            # 添加思考检查点（最后一个节点除外）
            if idx < len(node_list):
                result.append(f"[思考点{idx}] 请分析 {node_name} 的日志，识别问题并记录分析结果。")
                result.append(f"分析完成后，继续查看下一个节点（共{len(node_list)}个节点，当前第{idx}个）。")
                result.append("")
            else:
                result.append(f"[思考点{idx}] 请分析 {node_name} 的日志，识别问题并记录分析结果。")
                result.append("这是最后一个节点，分析完成后，请汇总所有节点的分析结果。")
                result.append("")
        
        result.append("[汇总提示] 所有节点已分析完毕，请汇总所有节点的分析结果，给出整体诊断和解决方案。")
        
        # 将结果写入文件
        result_text = "\n".join(result)
        try:
            # 确定result目录路径（相对于当前文件：rca/cl_agent/agent.py -> rca/result）
            current_dir = os.path.dirname(os.path.abspath(__file__))  # rca/cl_agent
            rca_dir = os.path.dirname(current_dir)  # rca
            result_dir = os.path.join(rca_dir, "result")
            
            # 确保目录存在
            os.makedirs(result_dir, exist_ok=True)
            
            # 生成带时间戳的文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"cluster_logs_{timestamp}.txt"
            file_path = os.path.join(result_dir, filename)
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(result_text)
            
            print(f"[INFO] 集群日志已保存到: {file_path}")
        except Exception as e:
            logging.warning(f"保存集群日志到文件失败: {e}")
            # 即使保存失败，也继续返回结果
        
        return result_text
    except Exception as e:
        return f"获取集群日志失败: {str(e)}"


@tool("get_node_log", description="获取指定节点的日志内容，用来分析指定节点的状态")
def get_node_log(node_name: str) -> str:
    """
    获取指定节点的日志内容。
    支持的节点：s1 NameNode, s1 DataNode, s2 DataNode, s3 DataNode, s3 SecondaryNameNode
    
    Args:
        node_name: 节点名称，可以是完整名称或简称
    
    Returns:
        原始日志文本，由Agent进行分析
    """
    print("调用get_node_log工具")
    try:
        log_content = get_node_log_by_name(node_name)
        return log_content
    except Exception as e:
        return f"获取节点日志失败: {str(e)}"


@tool("get_monitoring_metrics", description="获取集群的实时监控指标")
def get_monitoring_metrics() -> str:
    """
    获取HDFS集群的实时监控指标。
    返回NameNode和DataNodes的关键指标，包括节点状态、存储使用率、数据块状态等。
    
    Returns:
        格式化的监控指标文本
    """
    print("调用get_monitoring_metrics工具")
    try:
        from .monitor_collector import collect_all_metrics, format_metrics_for_display
        
        metrics = collect_all_metrics()
        html_content = format_metrics_for_display(metrics)
        
        # 将 HTML 转换为纯文本（移除 HTML 标签，保留内容）
        text_content = re.sub(r'<[^>]+>', '', html_content)
        text_content = text_content.replace('✅', '[正常]')
        text_content = text_content.replace('⚠️', '[异常]')
        text_content = text_content.replace('❌', '[错误]')
        try:
            # 确定metrics目录路径（相对于当前文件：rca/cl_agent/agent.py -> rca/metrics）
            current_dir = os.path.dirname(os.path.abspath(__file__))  # rca/cl_agent
            rca_dir = os.path.dirname(current_dir)  # rca
            result_dir = os.path.join(rca_dir, "metrics")
            
            # 确保目录存在
            os.makedirs(result_dir, exist_ok=True)
            
            # 生成带时间戳的文件名
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"metrics_{timestamp}.txt"
            file_path = os.path.join(result_dir, filename)
            
            # 写入文件
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            print(f"[INFO] 集群监控指标已保存到: {file_path}")
        except Exception as e:
            logging.warning(f"保存集群监控指-标到文件失败: {e}")
            # 即使保存失败，也继续返回结果
        return text_content
    except Exception as e:
        return f"获取监控指标失败: {str(e)}"


@tool("search_logs_by_keyword", description="在指定节点日志中搜索关键词，快速定位问题")
def search_logs_by_keyword(node_name: str, keyword: str, max_results: int = 50) -> str:
    """
    在指定节点日志中搜索关键词。
    
    Args:
        node_name: 节点名称（如：s1 NameNode, s2 DataNode）
        keyword: 搜索关键词（如：ERROR, WARN, Exception）
        max_results: 最大返回结果数（默认50）
    
    Returns:
        匹配的日志行及其上下文（时间戳、级别、消息）
    """
    print(f"调用search_logs_by_keyword工具: node={node_name}, keyword={keyword}, max_results={max_results}")
    try:
        # 1. 获取节点日志
        log_content = get_node_log_by_name(node_name)
        
        if not log_content or log_content.startswith("未找到节点") or log_content.startswith("读取日志失败"):
            return f"无法获取 {node_name} 的日志: {log_content}"
        
        # 2. 搜索关键词
        lines = log_content.split('\n')
        matches = []
        keyword_lower = keyword.lower()
        
        for i, line in enumerate(lines):
            if keyword_lower in line.lower():
                # 获取上下文（前后各2行）
                start = max(0, i - 2)
                end = min(len(lines), i + 3)
                context_lines = lines[start:end]
                context = '\n'.join(context_lines)
                matches.append({
                    'line_num': i + 1,
                    'context': context,
                    'matched_line': line
                })
                
                if len(matches) >= max_results:
                    break
        
        # 3. 格式化返回
        if not matches:
            return f"在 {node_name} 的日志中未找到关键词 '{keyword}'（共搜索了 {len(lines)} 行）"
        
        result = []
        result.append(f"在 {node_name} 的日志中找到 {len(matches)} 条匹配 '{keyword}' 的记录（共搜索 {len(lines)} 行）：\n")
        result.append("=" * 80)
        
        for idx, match in enumerate(matches, 1):
            result.append(f"\n[匹配 {idx}] 第 {match['line_num']} 行:")
            result.append("-" * 80)
            result.append(match['context'])
            result.append("-" * 80)
        
        if len(matches) >= max_results:
            result.append(f"\n[提示] 已显示前 {max_results} 条匹配结果，可能还有更多结果。")
        
        return "\n".join(result)
    except Exception as e:
        error_msg = f"搜索日志失败: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return error_msg


@tool("get_error_logs_summary", description="统计各节点的错误/警告数量，快速了解问题分布")
def get_error_logs_summary(node_name: Optional[str] = None) -> str:
    """
    统计各节点的错误/警告数量，按时间排序。
    
    Args:
        node_name: 节点名称（可选，不提供则统计所有节点）
    
    Returns:
        错误统计摘要（节点、错误数量、最新错误时间、错误类型分布）
    """
    print(f"调用get_error_logs_summary工具: node_name={node_name}")
    try:
        error_keywords = ['ERROR', 'WARN', 'Exception', 'FATAL', 'CRITICAL']
        all_summaries = []
        
        # 确定要统计的节点列表
        if node_name:
            # 只统计指定节点
            nodes_to_check = [node_name]
        else:
            # 统计所有节点
            nodes_to_check = [config["display_name"] for config in LOG_FILES_CONFIG]
        
        # 遍历节点统计错误
        for node in nodes_to_check:
            try:
                log_content = get_node_log_by_name(node)
                
                if not log_content or log_content.startswith("未找到节点") or log_content.startswith("读取日志失败"):
                    all_summaries.append({
                        'node': node,
                        'status': 'error',
                        'error': log_content
                    })
                    continue
                
                lines = log_content.split('\n')
                errors = []
                warnings = []
                
                # 解析日志，查找错误和警告
                for i, line in enumerate(lines):
                    line_upper = line.upper()
                    
                    # 检查是否是错误
                    if any(keyword in line_upper for keyword in ['ERROR', 'FATAL', 'CRITICAL', 'EXCEPTION']):
                        errors.append({
                            'line_num': i + 1,
                            'line': line.strip(),
                            'timestamp': _extract_timestamp(line)
                        })
                    # 检查是否是警告
                    elif 'WARN' in line_upper:
                        warnings.append({
                            'line_num': i + 1,
                            'line': line.strip(),
                            'timestamp': _extract_timestamp(line)
                        })
                
                # 统计错误类型
                error_types = {}
                for err in errors:
                    # 尝试提取错误类型（简单方法：取第一个单词或关键词）
                    error_type = "Unknown"
                    for keyword in ['ERROR', 'FATAL', 'CRITICAL', 'EXCEPTION']:
                        if keyword in err['line'].upper():
                            error_type = keyword
                            break
                    error_types[error_type] = error_types.get(error_type, 0) + 1
                
                all_summaries.append({
                    'node': node,
                    'status': 'success',
                    'total_lines': len(lines),
                    'error_count': len(errors),
                    'warning_count': len(warnings),
                    'error_types': error_types,
                    'latest_error': errors[-1] if errors else None,
                    'latest_warning': warnings[-1] if warnings else None,
                    'errors': errors[:10],  # 只保留前10个错误
                    'warnings': warnings[:10]  # 只保留前10个警告
                })
            except Exception as e:
                all_summaries.append({
                    'node': node,
                    'status': 'error',
                    'error': f"处理节点日志时出错: {str(e)}"
                })
        
        # 格式化返回结果
        result = []
        result.append("=" * 80)
        result.append("错误日志统计摘要")
        result.append("=" * 80)
        
        total_errors = 0
        total_warnings = 0
        
        for summary in all_summaries:
            if summary['status'] == 'error':
                result.append(f"\n[{summary['node']}] ❌ 无法获取日志: {summary['error']}")
                continue
            
            node = summary['node']
            error_count = summary['error_count']
            warning_count = summary['warning_count']
            total_errors += error_count
            total_warnings += warning_count
            
            result.append(f"\n[{node}]")
            result.append(f"  总行数: {summary['total_lines']}")
            result.append(f"  错误数: {error_count}")
            result.append(f"  警告数: {warning_count}")
            
            if error_count > 0:
                result.append(f"  错误类型分布: {summary['error_types']}")
                if summary['latest_error']:
                    result.append(f"  最新错误 (第{summary['latest_error']['line_num']}行): {summary['latest_error']['line'][:100]}")
            
            if warning_count > 0 and summary['latest_warning']:
                result.append(f"  最新警告 (第{summary['latest_warning']['line_num']}行): {summary['latest_warning']['line'][:100]}")
            
            # 显示前几个错误示例
            if summary['errors']:
                result.append(f"\n  错误示例（前{min(3, len(summary['errors']))}个）:")
                for err in summary['errors'][:3]:
                    result.append(f"    - 第{err['line_num']}行: {err['line'][:80]}")
        
        result.append("\n" + "=" * 80)
        result.append(f"总计: {len(nodes_to_check)} 个节点, {total_errors} 个错误, {total_warnings} 个警告")
        result.append("=" * 80)
        
        return "\n".join(result)
    except Exception as e:
        error_msg = f"统计错误日志失败: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return error_msg


def _extract_timestamp(line: str) -> Optional[str]:
    """从日志行中提取时间戳（简单实现）"""
    # 尝试匹配常见的时间戳格式
    # 例如: 2024-01-01 12:00:00, 2024-01-01T12:00:00, 01/01/2024 12:00:00
    patterns = [
        r'\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}',
        r'\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}',
        r'\d{2}/\d{2}/\d{4}\s+\d{2}:\d{2}:\d{2}',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, line)
        if match:
            return match.group(0)
    
    return None


# ==================== Hadoop集群操作工具 ====================

# Docker 容器白名单
ALLOWED_CONTAINERS = ['namenode', 'datanode1', 'datanode2']

# 允许的 Hadoop 操作类型
ALLOWED_HADOOP_OPERATIONS = {
    'start',      # 启动单个Hadoop服务或整个集群
    'stop',       # 停止单个Hadoop服务或整个集群
    'restart',    # 重启单个Hadoop服务（先stop再start）
}

# 容器到Hadoop服务类型的映射
CONTAINER_TO_DAEMON = {
    'namenode': 'namenode',
    'datanode1': 'datanode',
    'datanode2': 'datanode',
}

# Hadoop操作命令白名单（安全控制）
ALLOWED_HADOOP_COMMANDS = {
    # 查询类命令（只读，安全）
    'hdfs': ['dfsadmin', '-report', '-safemode', 'get', 'df', 'ls', 'du', 'fsck'],
    'hadoop': ['fs', 'version', 'classpath'],
    # 集群管理命令（需要谨慎使用）
    'start-dfs.sh': [],
    'stop-dfs.sh': [],
    'start-yarn.sh': [],
    'stop-yarn.sh': [],
    'start-all.sh': [],
    'stop-all.sh': [],
    # 节点管理命令
    'hadoop-daemon.sh': ['start', 'stop', 'status'],
    'hdfs-daemon.sh': ['start', 'stop', 'status'],
    'yarn-daemon.sh': ['start', 'stop', 'status'],
}


def _is_safe_hadoop_command(command: str) -> bool:
    """
    检查命令是否在允许的白名单中
    
    Args:
        command: 要执行的命令
    
    Returns:
        True表示命令安全，False表示命令不安全
    """
    command = command.strip()
    
    # 禁止的危险命令
    dangerous_keywords = ['rm', 'delete', 'format', 'kill', 'killall', 'shutdown', 'reboot']
    for keyword in dangerous_keywords:
        if keyword in command.lower():
            return False
    
    # 检查是否在允许的命令列表中
    parts = command.split()
    if not parts:
        return False
    
    base_cmd = parts[0]
    
    # 检查是否是hadoop/hdfs相关命令
    if base_cmd in ['hdfs', 'hadoop']:
        if len(parts) > 1:
            subcmd = parts[1]
            if base_cmd in ALLOWED_HADOOP_COMMANDS:
                allowed_subcmds = ALLOWED_HADOOP_COMMANDS[base_cmd]
                if not allowed_subcmds or subcmd in allowed_subcmds:
                    return True
    
    # 检查是否是脚本命令
    if base_cmd in ALLOWED_HADOOP_COMMANDS:
        return True
    
    # 检查是否是daemon命令
    if 'daemon.sh' in base_cmd:
        if len(parts) > 1:
            action = parts[1]
            if action in ['start', 'stop', 'status']:
                return True
    
    return False


def _get_docker_compose_path() -> str:
    """
    动态获取 docker-compose.yml 文件的绝对路径
    
    Returns:
        docker-compose.yml 的绝对路径
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))  # rca/cl_agent
    project_root = os.path.dirname(current_dir)  # rca
    docker_compose_file = os.path.join(project_root, "docker-compose.yml")
    return docker_compose_file


def _check_container_running(container: str) -> bool:
    """
    检查容器是否运行
    
    Args:
        container: 容器名称
    
    Returns:
        True表示容器运行中，False表示容器未运行
    """
    try:
        result = subprocess.run(
            ['docker', 'ps', '--format', '{{.Names}}'],
            shell=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            universal_newlines=True,
            timeout=5
        )
        if result.returncode == 0:
            running_containers = result.stdout.strip().split('\n')
            return container in running_containers
        return False
    except Exception:
        return False


def _validate_hadoop_operation_command(command_args: List[str]) -> Tuple[bool, Optional[str]]:
    """
    验证Hadoop操作命令参数的安全性
    
    Args:
        command_args: 命令参数列表（如 ["restart", "namenode"] 或 ["stop"]）
    
    Returns:
        (是否安全, 错误信息)，如果安全返回 (True, None)，否则返回 (False, 错误信息)
    """
    if not command_args or len(command_args) == 0:
        return False, "命令参数不能为空"
    
    # 获取操作类型（第一个参数）
    operation = command_args[0].lower()
    
    # 验证操作类型
    if operation not in ALLOWED_HADOOP_OPERATIONS:
        return False, f"不支持的操作类型 '{operation}'。支持的操作：{', '.join(sorted(ALLOWED_HADOOP_OPERATIONS))}"
    
    # 对于单节点操作（start/stop/restart + 容器名），验证容器名
    if len(command_args) >= 2:
        container = command_args[1]
        if container not in ALLOWED_CONTAINERS:
            return False, f"不支持的容器名称 '{container}'。支持的容器：{', '.join(ALLOWED_CONTAINERS)}"
    
    # 检查是否有危险参数
    dangerous_keywords = ['rm', 'delete', 'format', 'kill', 'killall', 'shutdown', 'reboot', 'down']
    for arg in command_args:
        if any(keyword in arg.lower() for keyword in dangerous_keywords):
            return False, f"检测到危险操作参数: {arg}"
    
    return True, None


@tool("hadoop_cluster_operation", description="执行Hadoop集群操作。command: stop/停止/关闭、start/启动、restart/重启。container: 可选，namenode/datanode1/datanode2。重要：如果用户要求关闭/启动/重启整个集群，container参数必须为None或不提供；只有操作单个节点时才指定container参数。")
def hadoop_cluster_operation(command: str, container: Optional[str] = None) -> str:
    """
    执行Hadoop集群操作（通过docker compose命令）。
    
    重要说明：
    - 启动整个集群：使用 docker compose up -d
    - 关闭整个集群：使用 docker compose stop
    - 对单个节点操作：使用 docker compose start/stop/restart 节点名称
    
    支持的操作：
    - "stop" 或 "停止" -> docker compose stop（整个集群）或 docker compose stop 节点名称（单个节点）
    - "start" 或 "启动" -> docker compose up -d（整个集群）或 docker compose start 节点名称（单个节点）
    - "restart" 或 "重启" -> docker compose restart 节点名称（仅支持单个节点）
    
    Args:
        command: 操作类型（"stop"/"停止"、"start"/"启动"、"restart"/"重启"）
        container: 节点名称（可选）。可选值：namenode, datanode1, datanode2
                   **重要**：
                   - 如果用户要求"关闭集群"、"停止集群"、"启动集群"等操作整个集群的命令，container必须为None或不提供
                   - 只有操作单个节点时（如"关闭datanode1"、"停止namenode"）才指定container参数
                   - 如果不指定container，会对整个集群执行操作
    
    Returns:
        操作执行结果
    """
    print(f"调用hadoop_cluster_operation工具: command={command}, container={container}")
    
    # 定义所有容器
    all_containers = ['namenode', 'datanode1', 'datanode2']
    
    # 验证容器名称（如果指定了）
    if container and container not in all_containers:
        return f"❌ 错误：不支持的容器 '{container}'。支持的容器：{', '.join(all_containers)}"
    
    # 解析命令
    command_lower = command.strip().lower()
    
    # 确定操作类型
    if any(x in command_lower for x in ['stop', '停止', '关闭']):
        operation_name = '停止'
    elif any(x in command_lower for x in ['start', '启动', '打开']):
        operation_name = '启动'
    elif any(x in command_lower for x in ['restart', '重启']):
        operation_name = '重启'
    else:
        return f"❌ 错误：不支持的操作 '{command}'。支持的操作：stop/停止、start/启动、restart/重启"
    
    # 获取 docker-compose.yml 文件所在目录（项目根目录）
    current_dir = os.path.dirname(os.path.abspath(__file__))  # rca/cl_agent
    project_root = os.path.dirname(current_dir)  # rca
    docker_compose_file = os.path.join(project_root, "docker-compose.yml")
    
    try:
        # 根据是否指定容器和操作类型，构建不同的命令
        if container:
            # 单个节点操作
            if operation_name == '重启':
                # restart 只支持单个节点
                docker_compose_cmd = ['docker', 'compose', '-f', docker_compose_file, 'restart', container]
                print(f"[{operation_name}] 节点 {container}: docker compose restart {container}")
            elif operation_name == '启动':
                docker_compose_cmd = ['docker', 'compose', '-f', docker_compose_file, 'start', container]
                print(f"[{operation_name}] 节点 {container}: docker compose start {container}")
            elif operation_name == '停止':
                docker_compose_cmd = ['docker', 'compose', '-f', docker_compose_file, 'stop', container]
                print(f"[{operation_name}] 节点 {container}: docker compose stop {container}")
        else:
            # 整个集群操作
            if operation_name == '启动':
                # 启动整个集群使用 docker compose up -d
                docker_compose_cmd = ['docker', 'compose', '-f', docker_compose_file, 'up', '-d']
                print(f"[{operation_name}] 整个集群: docker compose up -d")
            elif operation_name == '停止':
                # 停止整个集群使用 docker compose stop
                docker_compose_cmd = ['docker', 'compose', '-f', docker_compose_file, 'stop']
                print(f"[{operation_name}] 整个集群: docker compose stop")
            elif operation_name == '重启':
                # 重启整个集群不支持，需要指定节点
                return "❌ 错误：重启整个集群不支持，请指定具体的节点名称（container参数）"
        
        # 执行命令（在项目根目录执行）
        print(f"执行命令：{docker_compose_cmd}")
        result = subprocess.run(
            docker_compose_cmd,
            shell=False,
            cwd=project_root,  # 在项目根目录执行
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            universal_newlines=True,
            timeout=120  # 120秒超时（启动整个集群可能需要更长时间）
        )
        
        # 组合输出
        output_parts = []
        if container:
            output_parts.append(f"节点 {container} {operation_name}操作结果:")
        else:
            output_parts.append(f"整个集群{operation_name}操作结果:")
        output_parts.append("-" * 60)
        
        if result.stdout:
            output_parts.append(f"标准输出:\n{result.stdout}")
        if result.stderr:
            output_parts.append(f"标准错误:\n{result.stderr}")
        
        if result.returncode == 0:
            output_parts.append(f"\n✅ {operation_name}操作成功")
        else:
            output_parts.append(f"\n❌ {operation_name}操作失败 (返回码: {result.returncode})")
        
        output = "\n".join(output_parts)
        print(output)
        return output
        
    except subprocess.TimeoutExpired:
        error_msg = f"❌ 执行超时（超过120秒）"
        print(error_msg)
        return error_msg
    except Exception as e:
        error_msg = f"❌ 执行失败: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return error_msg
@tool("hadoop_auto_operation", description="""
执行Hadoop集群操作（在容器内启动/停止/重启Hadoop服务）。

**参数说明**：
- operation: 操作类型，可选值：start/stop/restart
- container: 容器名称（可选），可选值：namenode, datanode1, datanode2
  - 如果指定container，则操作单个节点的Hadoop服务
  - 如果不指定container（None），则操作整个集群（在namenode容器内执行start-dfs.sh/stop-dfs.sh）

**使用示例**：
- 启动NameNode: operation="start", container="namenode"
- 停止DataNode1: operation="stop", container="datanode1"
- 重启NameNode: operation="restart", container="namenode"
- 启动整个集群: operation="start", container=None
- 停止整个集群: operation="stop", container=None

**注意**：容器必须已运行，否则返回错误
""")
def hadoop_auto_operation(operation: str, container: Optional[str] = None) -> str:
    """
    执行Hadoop集群操作（在容器内执行Hadoop服务命令）。
    
    Args:
        operation: 操作类型（"start"/"stop"/"restart"）
        container: 容器名称（可选）。如果为None，则操作整个集群
    
    Returns:
        操作执行结果
    """
    print(f"调用hadoop_auto_operation工具: operation={operation}, container={container}")
    
    # 参数验证
    operation = operation.lower().strip()
    if operation not in ['start', 'stop', 'restart']:
        return f"❌ 错误：不支持的操作类型 '{operation}'。支持的操作：start, stop, restart"
    
    # 如果是单节点操作，验证容器名
    if container is not None:
        container = container.lower().strip()
        if container not in ALLOWED_CONTAINERS:
            return f"❌ 错误：不支持的容器名称 '{container}'。支持的容器：{', '.join(ALLOWED_CONTAINERS)}"
    
    # 构建command_args用于验证（兼容原有验证逻辑）
    if container:
        command_args = [operation, container]
    else:
        command_args = [operation]
    
    # 安全性验证
    is_safe, error_msg = _validate_hadoop_operation_command(command_args)
    if not is_safe:
        return f"❌ 错误：{error_msg}"
    
    # 确定是单节点操作还是整个集群操作
    is_cluster_operation = container is None
    
    # 确定操作名称（中文）
    operation_name_map = {
        'start': '启动',
        'stop': '停止',
        'restart': '重启'
    }
    operation_name = operation_name_map.get(operation, operation)
    
    # 如果是整个集群操作
    if is_cluster_operation:
        if operation == 'restart':
            return "❌ 错误：重启整个集群不支持，请先停止整个集群，再启动整个集群"
        
        # 检查namenode容器是否运行
        if not _check_container_running('namenode'):
            return "❌ 错误：namenode容器未运行，无法执行集群操作"
        
        # 构建命令：在namenode容器内执行start-dfs.sh或stop-dfs.sh
        script_name = 'start-dfs.sh' if operation == 'start' else 'stop-dfs.sh'
        # 使用su - hadoop切换到hadoop用户，并设置环境变量
        hadoop_env = "export HADOOP_HOME=/usr/local/hadoop && export PATH=$HADOOP_HOME/bin:$HADOOP_HOME/sbin:$PATH && "
        command_str = f"su - hadoop -c '{hadoop_env}{script_name}'"
        docker_command = ['docker', 'exec', 'namenode', 'sh', '-c', command_str]
        target_desc = "整个集群"
        
    else:
        # 单节点操作（container已经验证过，直接使用）
        
        # 检查容器是否运行
        if not _check_container_running(container):
            return f"❌ 错误：容器 '{container}' 未运行，无法执行操作"
        
        # 获取Hadoop服务类型
        daemon_type = CONTAINER_TO_DAEMON.get(container)
        if not daemon_type:
            return f"❌ 错误：无法确定容器 '{container}' 对应的Hadoop服务类型"
        
        # 处理重启操作（先stop再start）
        if operation == 'restart':
            # 先停止
            hadoop_env = "export HADOOP_HOME=/usr/local/hadoop && export PATH=$HADOOP_HOME/bin:$HADOOP_HOME/sbin:$PATH && "
            stop_command_str = f"su - hadoop -c '{hadoop_env}hdfs --daemon stop {daemon_type}'"
            stop_docker_command = ['docker', 'exec', container, 'sh', '-c', stop_command_str]
            
            # 再启动
            start_command_str = f"su - hadoop -c '{hadoop_env}hdfs --daemon start {daemon_type}'"
            start_docker_command = ['docker', 'exec', container, 'sh', '-c', start_command_str]
            
            try:
                # 执行停止
                print(f"执行停止命令：{' '.join(stop_docker_command)}")
                stop_result = subprocess.run(
                    stop_docker_command,
                    shell=False,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    universal_newlines=True,
                    timeout=30
                )
                
                # 等待一下
                import time
                time.sleep(2)
                
                # 执行启动
                print(f"执行启动命令：{' '.join(start_docker_command)}")
                start_result = subprocess.run(
                    start_docker_command,
                    shell=False,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    universal_newlines=True,
                    timeout=30
                )
                
                # 组合输出
                output_parts = []
                output_parts.append(f"节点 {container} 重启操作结果:")
                output_parts.append("-" * 60)
                output_parts.append("停止阶段:")
                if stop_result.stdout:
                    output_parts.append(f"标准输出:\n{stop_result.stdout}")
                if stop_result.stderr:
                    output_parts.append(f"标准错误:\n{stop_result.stderr}")
                output_parts.append("\n启动阶段:")
                if start_result.stdout:
                    output_parts.append(f"标准输出:\n{start_result.stdout}")
                if start_result.stderr:
                    output_parts.append(f"标准错误:\n{start_result.stderr}")
                
                if stop_result.returncode == 0 and start_result.returncode == 0:
                    output_parts.append(f"\n✅ 重启操作成功")
                    # 提示Agent调用验证工具
                    output_parts.append(f"\n💡 提示：建议调用 verify_cluster_health 工具验证修复是否成功")
                else:
                    output_parts.append(f"\n❌ 重启操作失败 (停止返回码: {stop_result.returncode}, 启动返回码: {start_result.returncode})")
                
                output = "\n".join(output_parts)
                print(output)
                return output
                
            except subprocess.TimeoutExpired:
                return "❌ 执行超时（超过30秒）"
            except Exception as e:
                return f"❌ 执行失败: {str(e)}"
        
        # 构建命令：在容器内执行hdfs --daemon命令
        hadoop_env = "export HADOOP_HOME=/usr/local/hadoop && export PATH=$HADOOP_HOME/bin:$HADOOP_HOME/sbin:$PATH && "
        command_str = f"su - hadoop -c '{hadoop_env}hdfs --daemon {operation} {daemon_type}'"
        docker_command = ['docker', 'exec', container, 'sh', '-c', command_str]
        target_desc = f"节点 {container}"
    
    # 执行命令
    try:
        print(f"执行命令：{' '.join(docker_command)}")
        result = subprocess.run(
            docker_command,
            shell=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            universal_newlines=True,
            timeout=60  # 60秒超时
        )
        
        # 组合输出
        output_parts = []
        output_parts.append(f"{target_desc} {operation_name}操作结果:")
        output_parts.append("-" * 60)
        
        if result.stdout:
            output_parts.append(f"标准输出:\n{result.stdout}")
        if result.stderr:
            output_parts.append(f"标准错误:\n{result.stderr}")
        
        if result.returncode == 0:
            output_parts.append(f"\n✅ {operation_name}操作成功")
        else:
            output_parts.append(f"\n❌ {operation_name}操作失败 (返回码: {result.returncode})")
        
        output = "\n".join(output_parts)
        print(output)
        return output
        
    except subprocess.TimeoutExpired:
        error_msg = f"❌ 执行超时（超过60秒）"
        print(error_msg)
        return error_msg
    except Exception as e:
        error_msg = f"❌ 执行失败: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return error_msg






@tool("execute_hadoop_command", description="""
执行Hadoop命令（hdfs、hadoop等）。在容器内执行Hadoop管理命令。

**参数格式**：
- command_args: List[str]，Hadoop命令参数列表
- 示例：["hdfs", "dfsadmin", "-report"]、["hdfs", "dfsadmin", "-safemode", "get"]

**支持的命令**：hdfs dfsadmin、hdfs fsck、hdfs dfs等（只读查询命令）
**执行容器**：通常为namenode（集群级命令）
""")
def execute_hadoop_command(command_args: List[str]) -> str:
    """
    执行Hadoop命令（在容器内执行）。
    
    Args:
        command_args: Hadoop命令参数列表，例如：
            - ["hdfs", "dfsadmin", "-report"] - 查看集群状态
            - ["hdfs", "dfsadmin", "-safemode", "get"] - 查看安全模式
            - ["hdfs", "fsck", "/"] - 检查文件系统
    
    Returns:
        命令执行结果
    """
    print(f"调用execute_hadoop_command工具: command_args={command_args}")
    
    # 参数验证
    if not isinstance(command_args, list):
        return f"❌ 错误：command_args 必须是列表格式，当前类型：{type(command_args).__name__}"
    
    if len(command_args) == 0:
        return "❌ 错误：command_args 不能为空"
    
    # 构建完整命令
    base_cmd = command_args[0]
    if base_cmd not in ['hdfs', 'hadoop']:
        return f"❌ 错误：不支持的命令 '{base_cmd}'。支持的命令：hdfs, hadoop"
    
    # 安全检查
    full_command_str = ' '.join(command_args)
    if not _is_safe_hadoop_command(full_command_str):
        return f"❌ 错误：命令 '{full_command_str}' 不在允许的白名单中或包含危险操作"
    
    # 确定目标容器（集群级命令通常在namenode执行）
    target_container = "namenode"
    
    # 构建docker exec命令
    # 设置Hadoop环境变量并切换到hadoop用户
    hadoop_env = "export HADOOP_HOME=/usr/local/hadoop && export PATH=$HADOOP_HOME/bin:$HADOOP_HOME/sbin:$PATH && "
    command_str = f"su - hadoop -c '{hadoop_env}{' '.join(command_args)}'"
    full_command = ['docker', 'exec', target_container, 'sh', '-c', command_str]
    
    try:
        # 执行命令
        print(f"执行命令：{' '.join(full_command)}")
        result = subprocess.run(
            full_command,
            shell=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            universal_newlines=True,
            timeout=60  # 60秒超时
        )
        
        # 组合输出
        output_parts = []
        output_parts.append(f"Hadoop命令执行结果:")
        output_parts.append("-" * 60)
        
        if result.stdout:
            output_parts.append(f"标准输出:\n{result.stdout}")
        if result.stderr:
            output_parts.append(f"标准错误:\n{result.stderr}")
        
        if result.returncode == 0:
            output_parts.append(f"\n✅ 命令执行成功")
        else:
            output_parts.append(f"\n❌ 命令执行失败 (返回码: {result.returncode})")
        
        output = "\n".join(output_parts)
        print(output)
        return output
        
    except subprocess.TimeoutExpired:
        error_msg = f"❌ 执行超时（超过60秒）"
        print(error_msg)
        return error_msg
    except Exception as e:
        error_msg = f"❌ 执行失败: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return error_msg


# ==================== 验证配置 ====================

# 验证配置常量
VERIFY_CONFIG = {
    "expected_datanode_count": 3,  # 期望的DataNode数量
    "max_missing_blocks": 0,  # 允许的最大缺失数据块数
    "max_corrupt_blocks": 0,  # 允许的最大损坏数据块数
    "check_error_logs": False,  # 是否检查错误日志（暂时关闭，避免过于复杂）
    "error_log_time_window": 300  # 检查最近N秒的错误日志
}


@tool("verify_cluster_health", description="""
验证HDFS集群健康状态。在修复操作后调用此工具，检查集群是否恢复正常。

**功能**：
- 检查所有DataNode是否在线（期望3个DataNode）
- 检查NameNode状态是否正常
- 检查数据块状态（是否有缺失或损坏的数据块）
- 返回结构化的验证结果

**使用场景**：
- 执行修复操作（如重启节点、启动集群）后，调用此工具验证修复是否成功
- 定期检查集群健康状态
- 故障诊断后验证问题是否已解决

**返回结果**：
- overall_status: success/failed/partial（整体状态）
- 各项检查的详细结果
- 修复建议（如果发现问题）
""")
def verify_cluster_health() -> str:
    """
    验证HDFS集群健康状态。
    
    检查项：
    1. DataNode状态：检查所有3个DataNode是否在线
    2. NameNode状态：检查NameNode是否正常
    3. 数据块状态：检查是否有缺失或损坏的数据块
    
    Returns:
        结构化的验证结果字符串
    """
    print("调用verify_cluster_health工具")
    
    try:
        # 1. 获取监控指标
        from .monitor_collector import collect_all_metrics
        metrics = collect_all_metrics()
        
        # 2. 获取集群报告
        cluster_report_result = execute_hadoop_command(["hdfs", "dfsadmin", "-report"])
        
        # 3. 执行验证检查
        checks = {}
        overall_status = "success"
        
        # 检查1：DataNode状态
        namenode_metrics = metrics.get("namenode", {})
        datanodes_metrics = metrics.get("datanodes", {})
        
        # 统计在线DataNode数量
        online_datanodes = 0
        offline_datanodes = []
        
        for node_name, node_data in datanodes_metrics.items():
            if (node_data.get("status") != "error" and 
                node_data.get("metrics", {}).get("datanode_status", {}).get("value") == "running"):
                online_datanodes += 1
            else:
                offline_datanodes.append(node_name)
        
        expected_count = VERIFY_CONFIG["expected_datanode_count"]
        datanode_check_passed = online_datanodes >= expected_count
        
        checks["datanode_status"] = {
            "status": "pass" if datanode_check_passed else "fail",
            "details": f"在线DataNode: {online_datanodes}/{expected_count}",
            "offline_nodes": offline_datanodes if offline_datanodes else None
        }
        
        if not datanode_check_passed:
            overall_status = "failed"
        
        # 检查2：NameNode状态
        namenode_status = namenode_metrics.get("status", "error")
        namenode_check_passed = namenode_status == "normal"
        
        checks["namenode_status"] = {
            "status": "pass" if namenode_check_passed else "fail",
            "details": f"NameNode状态: {namenode_status}"
        }
        
        if not namenode_check_passed:
            overall_status = "failed"
        
        # 检查3：数据块状态（从集群报告中解析）
        missing_blocks = 0
        corrupt_blocks = 0
        
        # 从集群报告中提取数据块信息
        if "标准输出" in cluster_report_result:
            report_lines = cluster_report_result.split("\n")
            for line in report_lines:
                if "Missing blocks:" in line:
                    # 尝试提取数字
                    match = re.search(r'Missing blocks:\s*(\d+)', line)
                    if match:
                        missing_blocks = int(match.group(1))
                elif "Blocks with corrupt replicas:" in line:
                    match = re.search(r'Blocks with corrupt replicas:\s*(\d+)', line)
                    if match:
                        corrupt_blocks = int(match.group(1))
        
        block_check_passed = (missing_blocks <= VERIFY_CONFIG["max_missing_blocks"] and 
                             corrupt_blocks <= VERIFY_CONFIG["max_corrupt_blocks"])
        
        checks["block_status"] = {
            "status": "pass" if block_check_passed else "fail",
            "details": f"缺失数据块: {missing_blocks}, 损坏数据块: {corrupt_blocks}",
            "missing_blocks": missing_blocks,
            "corrupt_blocks": corrupt_blocks
        }
        
        if not block_check_passed:
            if overall_status == "success":
                overall_status = "partial"  # 如果其他都正常，但数据块有问题，算部分成功
            else:
                overall_status = "failed"
        
        # 检查4：活跃DataNode数量（从NameNode指标）
        live_datanodes = 0
        dead_datanodes = 0
        
        if namenode_metrics.get("status") == "normal":
            live_datanodes_metric = namenode_metrics.get("metrics", {}).get("live_datanodes", {})
            dead_datanodes_metric = namenode_metrics.get("metrics", {}).get("dead_datanodes", {})
            
            # 尝试从heartbeat_value获取
            live_datanodes = live_datanodes_metric.get("heartbeat_value", 0)
            dead_datanodes = dead_datanodes_metric.get("heartbeat_value", 0)
        
        heartbeat_check_passed = (live_datanodes >= expected_count and dead_datanodes == 0)
        
        checks["heartbeat_status"] = {
            "status": "pass" if heartbeat_check_passed else "fail",
            "details": f"活跃DataNode: {live_datanodes}, 死掉的DataNode: {dead_datanodes}"
        }
        
        if not heartbeat_check_passed:
            if overall_status == "success":
                overall_status = "partial"
            else:
                overall_status = "failed"
        
        # 生成验证结果摘要
        result_parts = []
        result_parts.append("=" * 60)
        result_parts.append("集群健康状态验证结果")
        result_parts.append("=" * 60)
        result_parts.append(f"整体状态: {overall_status.upper()}")
        result_parts.append("")
        
        # 详细检查结果
        result_parts.append("检查项详情:")
        result_parts.append("-" * 60)
        
        for check_name, check_result in checks.items():
            status_icon = "✅" if check_result["status"] == "pass" else "❌"
            check_name_cn = {
                "datanode_status": "DataNode状态",
                "namenode_status": "NameNode状态",
                "block_status": "数据块状态",
                "heartbeat_status": "心跳状态"
            }.get(check_name, check_name)
            
            result_parts.append(f"{status_icon} {check_name_cn}: {check_result['details']}")
            
            # 如果有离线节点，显示详细信息
            if check_result.get("offline_nodes"):
                result_parts.append(f"   离线节点: {', '.join(check_result['offline_nodes'])}")
        
        result_parts.append("")
        
        # 生成修复建议
        if overall_status != "success":
            result_parts.append("修复建议:")
            result_parts.append("-" * 60)
            
            if not datanode_check_passed:
                result_parts.append(f"⚠️ 发现 {len(offline_datanodes)} 个DataNode离线: {', '.join(offline_datanodes)}")
                result_parts.append("   建议：检查离线节点的容器状态，尝试重启对应的Hadoop服务")
            
            if not namenode_check_passed:
                result_parts.append("⚠️ NameNode状态异常")
                result_parts.append("   建议：检查NameNode日志，尝试重启NameNode服务")
            
            if not block_check_passed:
                result_parts.append(f"⚠️ 发现数据块问题：缺失 {missing_blocks} 个，损坏 {corrupt_blocks} 个")
                result_parts.append("   建议：检查DataNode日志，可能需要重新复制数据块")
            
            if not heartbeat_check_passed:
                result_parts.append(f"⚠️ 心跳异常：{dead_datanodes} 个DataNode未向NameNode发送心跳")
                result_parts.append("   建议：检查网络连接和DataNode服务状态")
        else:
            result_parts.append("✅ 所有检查项通过，集群状态正常")
        
        result_parts.append("")
        result_parts.append("=" * 60)
        
        result = "\n".join(result_parts)
        print(result)
        return result
        
    except Exception as e:
        error_msg = f"❌ 验证过程发生错误: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return error_msg


# ==================== Agent创建 ====================

def create_agent_instance(model_name: str = "qwen-8b"):
    """
    根据模型名称创建Agent实例
    
    Args:
        model_name: 模型名称，可选值：
            - "qwen-8b": Qwen-8B (vLLM本地部署)
            - "gpt-4o": GPT-4o (第三方API)
            - "deepseek-r1": DeepSeek-R1 (第三方API)
    
    Returns:
        Agent 实例
    """
    print(f"[DEBUG] 开始创建Agent实例 - 模型: {model_name}")
    
    # 导入知识库模块（暂时注释，不使用知识库）
    # try:
    #     from .knowledge_base import search_operation_knowledge
    # except ImportError:
    #     from lc_agent.knowledge_base import search_operation_knowledge
    
    # 定义操作知识库检索工具（暂时注释，不使用知识库）
    # @tool("search_operation_knowledge", description="检索Hadoop集群操作知识库，获取操作命令格式、示例和说明。如果不确定如何执行某个操作，请先使用此工具检索。")
    # def search_operation_knowledge_tool(query: str) -> str:
    #     """检索操作知识库"""
    #     return search_operation_knowledge(query)
    
    llm = create_llm(model_name)
    tools = [get_cluster_logs, get_node_log, get_monitoring_metrics, website_search, hadoop_auto_operation, execute_hadoop_command, verify_cluster_health]
 
    
    system_prompt = """你是HDFS集群问题诊断专家。

重要提示：
- 工具调用后，基于返回结果进行分析
- 工具返回错误时，说明原因并建议解决方案

**分批分析要求**：
调用get_cluster_logs时，按思考检查点逐个分析节点日志，最后汇总。

**集群操作工具使用说明**：
- hadoop_auto_operation 用于在容器内执行Hadoop服务操作（启动/停止/重启节点或集群）
- execute_hadoop_command 用于执行Hadoop管理命令（查看集群状态、安全模式等）
- verify_cluster_health 用于验证集群健康状态（修复操作后必须调用此工具验证修复是否成功）

**修复操作流程（重要）**：
1. 执行修复操作（使用 hadoop_auto_operation 或 execute_hadoop_command）
2. 修复操作成功后，**必须**调用 verify_cluster_health 工具验证修复是否成功
3. 根据验证结果判断：
   - 如果验证通过（overall_status=success），说明修复成功
   - 如果验证失败（overall_status=failed），说明修复未成功，需要进一步诊断
   - 如果验证部分成功（overall_status=partial），说明部分问题已解决，但仍有问题需要处理

请用专业、清晰的语言回答。"""
    
    print(f"[DEBUG] 正在创建Agent（使用 {model_name} 模型）...")
    agent = create_agent(
        model=llm,
        tools=tools,
        system_prompt=system_prompt
    )
    
    print(f"[DEBUG] ✅ Agent实例创建成功 - 模型: {model_name}")
    
    return agent

# ==================== 文档导出功能 ====================

def parse_markdown_to_word_paragraphs(doc, text: str):
    """
    将Markdown格式的文本解析并添加到Word文档中
    
    Args:
        doc: python-docx Document对象
        text: Markdown格式的文本
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    lines = text.split('\n')
    i = 0
    in_list = False
    list_level = 0
    current_para_lines = []  # 用于合并多行段落
    
    while i < len(lines):
        line = lines[i].rstrip()  # 只移除右侧空格，保留左侧缩进信息
        
        # 处理分隔线（---或===）
        if line.strip().startswith('---') or (line.strip().startswith('===') and len(line.strip()) >= 3):
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if in_list:
                in_list = False
                list_level = 0
            doc.add_paragraph()  # 空行作为分隔
            i += 1
            continue
        
        # 处理标题（###、####等）
        if line.strip().startswith('#'):
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if in_list:
                in_list = False
                list_level = 0
            
            # 计算标题级别
            stripped = line.strip()
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            
            # 提取标题文本（移除#号和空格）
            title_text = stripped[level:].strip()
            
            # 移除Markdown格式标记（**、`等），但保留文本
            title_text = re.sub(r'\*\*(.*?)\*\*', r'\1', title_text)  # 移除加粗，保留文本
            title_text = re.sub(r'`(.*?)`', r'\1', title_text)  # 移除代码标记，保留文本
            
            # 添加标题（Word标题级别：1-9，我们限制在1-3）
            heading_level = min(level, 3)
            doc.add_heading(title_text, heading_level)
            i += 1
            continue
        
        # 处理有序列表（1.、2.等）
        ordered_match = re.match(r'^(\d+)[\.、]\s+(.+)$', line.strip())
        if ordered_match:
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if not in_list:
                in_list = True
                list_level = 0
            
            item_text = ordered_match.group(2)
            # 处理列表项中的格式（保留格式，不转换为纯文本）
            para = doc.add_paragraph(style='List Number')
            _add_formatted_text(para, item_text)
            i += 1
            continue
        
        # 处理无序列表（-、*等，但排除分隔线）
        if line.strip() and (line.strip().startswith('- ') or line.strip().startswith('* ')):
            if current_para_lines:
                # 先处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if not in_list:
                in_list = True
                list_level = 0
            
            # 提取列表项文本（移除 - 或 * 和空格）
            item_text = line.strip()[2:].strip()
            # 处理列表项中的格式（保留格式）
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, item_text)
            i += 1
            continue
        
        # 处理空行
        if not line.strip():
            if current_para_lines:
                # 处理累积的段落
                para = doc.add_paragraph()
                _add_formatted_text(para, ' '.join(current_para_lines))
                current_para_lines = []
            if in_list:
                in_list = False
                list_level = 0
            # 添加空行
            doc.add_paragraph()
            i += 1
            continue
        
        # 处理普通文本行
        if in_list:
            in_list = False
            list_level = 0
        
        # 累积到当前段落（处理多行段落）
        current_para_lines.append(line.strip())
        i += 1
    
    # 处理最后剩余的段落
    if current_para_lines:
        para = doc.add_paragraph()
        _add_formatted_text(para, ' '.join(current_para_lines))


def _process_inline_formatting(text: str) -> str:
    """处理行内格式标记，返回纯文本"""
    # 移除加粗标记
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # 移除代码标记
    text = re.sub(r'`(.*?)`', r'\1', text)
    # 移除斜体标记
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text


def _add_formatted_text(para, text: str):
    """
    将包含Markdown格式的文本添加到段落中，保留格式
    
    Args:
        para: python-docx Paragraph对象
        text: 包含Markdown格式的文本
    """
    from docx.shared import Pt, RGBColor
    
    if not text:
        return
    
    # 使用正则表达式匹配所有格式标记（加粗、代码、斜体）
    # 按顺序处理，避免嵌套问题
    pattern = r'(\*\*.*?\*\*|`.*?`|\*.*?\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # 处理加粗文本 **text**
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            bold_text = part[2:-2]  # 移除 **
            # 检查加粗文本内部是否还有代码标记
            if '`' in bold_text:
                # 递归处理
                _add_formatted_text(para, bold_text)
            else:
                run = para.add_run(bold_text)
                run.bold = True
        
        # 处理代码文本 `text`
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            code_text = part[1:-1]  # 移除 `
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        # 处理斜体文本 *text*（但不是加粗）
        elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
            italic_text = part[1:-1]  # 移除 *
            run = para.add_run(italic_text)
            run.italic = True
        
        # 普通文本
        else:
            # 如果还有未处理的格式标记，递归处理
            if '**' in part or '`' in part or (part.startswith('*') and part.endswith('*') and len(part) > 2):
                _add_formatted_text(para, part)
            else:
                para.add_run(part)


def export_to_word(analysis_result: str, output_path: str) -> str:
    """
    将Agent的分析结果导出为Word文档
    
    Args:
        analysis_result: Agent返回的分析结果文本
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        
        doc = Document()
        
        # 标题
        title = doc.add_heading('HDFS集群问题诊断报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 时间戳
        timestamp = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
        time_para = doc.add_paragraph(f'生成时间: {timestamp}')
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # 空行
        
        # 执行摘要
        doc.add_heading('执行摘要', 1)
        summary_para = doc.add_paragraph('本报告基于HDFS集群监控数据和日志分析生成，包含监控指标、问题诊断、原因分析和解决方案。')
        
        doc.add_paragraph()  # 空行
        
        # 添加集群监控指标表格
        doc.add_heading('集群监控指标', 1)
        
        try:
            from .monitor_collector import collect_all_metrics
            metrics_data = collect_all_metrics()
            
            # NameNode 指标表格
            doc.add_heading('NameNode 监控指标', 2)
            nn_table = doc.add_table(rows=1, cols=3)
            nn_table.style = 'Light Grid Accent 1'
            
            # 表头
            header_cells = nn_table.rows[0].cells
            header_cells[0].text = '指标名称'
            header_cells[1].text = '指标值'
            header_cells[2].text = '状态'
            
            # 设置表头样式
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # 添加NameNode指标数据
            if metrics_data["namenode"].get("status") != "error":
                for key, metric in metrics_data["namenode"].get("metrics", {}).items():
                    row_cells = nn_table.add_row().cells
                    row_cells[0].text = metric['name']
                    row_cells[1].text = str(metric['value'])
                    status_text = "正常" if metric["status"] == "normal" else "异常"
                    row_cells[2].text = status_text
                    
                    # 设置状态列颜色
                    if metric["status"] == "normal":
                        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                    else:
                        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # 红色
            else:
                row_cells = nn_table.add_row().cells
                row_cells[0].text = "错误"
                row_cells[1].text = metrics_data['namenode'].get('error', '未知错误')
                row_cells[2].text = "异常"
                row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            doc.add_paragraph()  # 空行
            
            # DataNode 指标表格
            doc.add_heading('DataNode 监控指标', 2)
            for node_name, node_data in metrics_data["datanodes"].items():
                doc.add_heading(f'DataNode {node_name.upper()}', 3)
                dn_table = doc.add_table(rows=1, cols=3)
                dn_table.style = 'Light Grid Accent 1'
                
                # 表头
                header_cells = dn_table.rows[0].cells
                header_cells[0].text = '指标名称'
                header_cells[1].text = '指标值'
                header_cells[2].text = '状态'
                
                # 设置表头样式
                for cell in header_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                
                # 添加DataNode指标数据
                if node_data.get("status") != "error":
                    for key, metric in node_data.get("metrics", {}).items():
                        row_cells = dn_table.add_row().cells
                        # 移除指标名称中的节点前缀，使表格更简洁
                        metric_name = metric['name'].replace(f'DataNode {node_name} ', '')
                        row_cells[0].text = metric_name
                        row_cells[1].text = str(metric['value'])
                        status_text = "正常" if metric["status"] == "normal" else "异常"
                        row_cells[2].text = status_text
                        
                        # 设置状态列颜色
                        if metric["status"] == "normal":
                            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        else:
                            row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                else:
                    row_cells = dn_table.add_row().cells
                    row_cells[0].text = "错误"
                    row_cells[1].text = node_data.get('error', '未知错误')
                    row_cells[2].text = "异常"
                    row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                
                doc.add_paragraph()  # 空行
            
            # 添加监控数据时间戳
            doc.add_paragraph(f'监控数据采集时间: {metrics_data["timestamp"]}')
            doc.add_paragraph()  # 空行
            
        except Exception as e:
            doc.add_paragraph(f'⚠️ 获取监控指标失败: {str(e)}')
            doc.add_paragraph()  # 空行
        
        doc.add_paragraph()  # 空行
        
        # 问题分析
        doc.add_heading('问题分析', 1)
        
        # 使用Markdown解析函数处理分析结果
        parse_markdown_to_word_paragraphs(doc, analysis_result)
        
        doc.add_paragraph()  # 空行
        
        # 解决方案
        doc.add_heading('解决方案', 1)
        solution_para = doc.add_paragraph('根据上述分析，建议采取以下措施：')
        solution_para.add_run('\n\n请参考问题分析部分的详细说明。')
        
        doc.add_paragraph()  # 空行
        
        # 预防措施
        doc.add_heading('预防措施', 1)
        prevention_para = doc.add_paragraph('为避免类似问题再次发生，建议：')
        prevention_para.add_run('\n1. 定期检查集群监控指标\n2. 及时处理异常日志\n3. 建立完善的监控告警机制')
        
        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    except Exception as e:
        raise Exception(f"导出Word文档失败: {str(e)}")


def export_to_pdf(analysis_result: str, output_path: str) -> str:
    """
    将Agent的分析结果导出为PDF文档
    
    Args:
        analysis_result: Agent返回的分析结果文本
        output_path: 输出文件路径
    
    Returns:
        输出文件路径
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.colors import black
        
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=black,
            spaceAfter=30,
            alignment=1  # 居中
        )
        
        # 小标题样式
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=black,
            spaceAfter=12,
            spaceBefore=12
        )
        
        # 正文样式
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=6
        )
        
        # 标题
        story.append(Paragraph('HDFS集群问题诊断报告', title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 时间戳
        timestamp = datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')
        story.append(Paragraph(f'生成时间: {timestamp}', styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # 执行摘要
        story.append(Paragraph('执行摘要', heading_style))
        story.append(Paragraph('本报告基于HDFS集群监控数据和日志分析生成，包含监控指标、问题诊断、原因分析和解决方案。', normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 添加集群监控指标表格
        story.append(Paragraph('集群监控指标', heading_style))
        
        try:
            from .monitor_collector import collect_all_metrics
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            
            metrics_data = collect_all_metrics()
            
            # NameNode 指标表格
            story.append(Paragraph('NameNode 监控指标', heading_style))
            
            if metrics_data["namenode"].get("status") != "error":
                # 准备表格数据和样式
                nn_data = [['指标名称', '指标值', '状态']]
                table_styles = [
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ]
                
                row_idx = 1
                for key, metric in metrics_data["namenode"].get("metrics", {}).items():
                    status_text = "正常" if metric["status"] == "normal" else "异常"
                    nn_data.append([metric['name'], str(metric['value']), status_text])
                    
                    # 设置状态列颜色
                    status_color = colors.green if metric["status"] == "normal" else colors.red
                    table_styles.append(('TEXTCOLOR', (2, row_idx), (2, row_idx), status_color))
                    row_idx += 1
                
                # 创建表格
                nn_table = Table(nn_data, colWidths=[2.5*inch, 2*inch, 1*inch])
                nn_table.setStyle(TableStyle(table_styles))
                
                story.append(nn_table)
                story.append(Spacer(1, 0.2*inch))
            else:
                story.append(Paragraph(f'错误: {metrics_data["namenode"].get("error", "未知错误")}', normal_style))
                story.append(Spacer(1, 0.2*inch))
            
            # DataNode 指标表格
            story.append(Paragraph('DataNode 监控指标', heading_style))
            
            for node_name, node_data in metrics_data["datanodes"].items():
                story.append(Paragraph(f'DataNode {node_name.upper()}', heading_style))
                
                if node_data.get("status") != "error":
                    # 准备表格数据和样式
                    dn_data = [['指标名称', '指标值', '状态']]
                    table_styles = [
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTSIZE', (0, 1), (-1, -1), 9),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    ]
                    
                    row_idx = 1
                    for key, metric in node_data.get("metrics", {}).items():
                        metric_name = metric['name'].replace(f'DataNode {node_name} ', '')
                        status_text = "正常" if metric["status"] == "normal" else "异常"
                        dn_data.append([metric_name, str(metric['value']), status_text])
                        
                        # 设置状态列颜色
                        status_color = colors.green if metric["status"] == "normal" else colors.red
                        table_styles.append(('TEXTCOLOR', (2, row_idx), (2, row_idx), status_color))
                        row_idx += 1
                    
                    # 创建表格
                    dn_table = Table(dn_data, colWidths=[2.5*inch, 2*inch, 1*inch])
                    dn_table.setStyle(TableStyle(table_styles))
                    
                    story.append(dn_table)
                else:
                    story.append(Paragraph(f'错误: {node_data.get("error", "未知错误")}', normal_style))
                
                story.append(Spacer(1, 0.2*inch))
            
            # 添加监控数据时间戳
            story.append(Paragraph(f'监控数据采集时间: {metrics_data["timestamp"]}', normal_style))
            story.append(Spacer(1, 0.2*inch))
            
        except Exception as e:
            story.append(Paragraph(f'⚠️ 获取监控指标失败: {str(e)}', normal_style))
            story.append(Spacer(1, 0.2*inch))
        
        # 问题分析
        story.append(Paragraph('问题分析', heading_style))
        
        # 将分析结果分段处理
        paragraphs = analysis_result.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 转义HTML特殊字符
                para_text = para_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # 检查是否是标题
                if re.match(r'^[0-9]+[\.、]', para_text) or any(keyword in para_text for keyword in ['问题', '原因', '解决', '建议']):
                    if len(para_text) < 50:
                        story.append(Paragraph(para_text, heading_style))
                    else:
                        story.append(Paragraph(para_text, normal_style))
                else:
                    story.append(Paragraph(para_text, normal_style))
        
        story.append(Spacer(1, 0.2*inch))
         
        # 解决方案
        story.append(Paragraph('解决方案', heading_style))
        story.append(Paragraph('根据上述分析，建议采取以下措施：', normal_style))
        story.append(Paragraph('请参考问题分析部分的详细说明。', normal_style))
        story.append(Spacer(1, 0.2*inch))
        
        # 预防措施
        story.append(Paragraph('预防措施', heading_style))
        story.append(Paragraph('为避免类似问题再次发生，建议：', normal_style))
        story.append(Paragraph('1. 定期检查集群监控指标', normal_style))
        story.append(Paragraph('2. 及时处理异常日志', normal_style))
        story.append(Paragraph('3. 建立完善的监控告警机制', normal_style))
        
        # 生成PDF
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.build(story)
        
        return output_path
    except ImportError:
        raise ImportError("需要安装 reportlab: pip install reportlab")
    except Exception as e:
        raise Exception(f"导出PDF文档失败: {str(e)}")

# ==================== 主函数（用于测试） ====================

if __name__ == "__main__":
    # 测试Agent创建
    print("[INFO] 正在创建Agent...")
    agent = create_agent_instance()
    print("[INFO] Agent创建成功！")
    
    # 测试调用
    config = {"configurable": {"thread_id": "test_1"}}
    result = agent.invoke(
        {"messages": [{"role": "user", "content": "查看集群状态"}]},
        config=config
    )
    
    print("\n[结果]")
    print(result['messages'][-1].content)

